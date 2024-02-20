import base64
import io
import json
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, timedelta
import pytz
from django.db import transaction, IntegrityError
from minio import Minio
from io import BytesIO
from dotenv import load_dotenv
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from uuid import uuid4
from docx.shared import Inches
from docx.shared import Pt
from staffing_table.models import StaffingTable, Vacancy
from birth_info.models import BirthInfo
from decree.models import DecreeList, TransferInfo, RankUpInfo, AppointmentInfo, OtpuskInfo, KomandirovkaInfo, Base, \
    FiringInfo
from education.models import Education, AcademicDegree
from education.serializers import EducationSerializer, AcademicDegreeSerializer
from location.models import Department
from military_rank.models import RankInfo, MilitaryRank
from person.models import Person, Vacation, Holidays
from photo.models import Photo
from position.models import Position, PositionInfo
from working_history.models import WorkingHistory
from military_rank.tasks import create_rank_info_after_months
import os

load_dotenv()

MINIO_ENDPOINT = os.getenv('MINIO_ENDPOINT')
MINIO_ACCESS_KEY = os.getenv('MINIO_ACCESS_KEY')
MINIO_SECRET_KEY = os.getenv('MINIO_SECRET_KEY')
MINIO_SECURE = os.getenv('MINIO_SECURE') == 'True'
MINIO_BUCKET_NAME = os.getenv('MINIO_BUCKET_NAME')


@csrf_exempt
def generate_work_reference(request, person_id):
    # Fetch the necessary data
    person = get_object_or_404(Person, pk=person_id)
    birth_info = BirthInfo.objects.get(personId=person)
    birth_date = str(birth_info.birth_date)
    birth_date_format = datetime.strptime(birth_date, '%Y-%m-%d')
    formatted_date = birth_date_format.strftime('%d.%m.%Y')

    try:
        rankInfo = RankInfo.objects.get(person=person)
    except RankInfo.DoesNotExist:
        rankInfo = None

    education_objects = Education.objects.filter(personId=person.id)
    education_data = EducationSerializer(education_objects, many=True).data

    if len(education_data) != 0:
        first_education = education_data[0]
        date_edu_string = first_education['educationDateOut']  # Assuming 'first_education' is an OrderedDict
        date_obj = datetime.strptime(date_edu_string, '%Y-%m-%d')

    academic_degrees_objects = AcademicDegree.objects.filter(personId=person.id)
    academic_degrees_data = AcademicDegreeSerializer(academic_degrees_objects, many=True).data
    if len(academic_degrees_data) != 0:
        first_academic_degree = academic_degrees_data[0]
        date_academ_string = first_academic_degree['academicDiplomaDate']
        date_academ_obj = datetime.strptime(date_academ_string, '%Y-%m-%d')

    persons_photo = Photo.objects.get(personId=person)
    photo_base64 = persons_photo.photoBinary

    # Replace with your photo field
    photo_binary = base64.b64decode(photo_base64)
    image = io.BytesIO(photo_binary)

    def calculate_experience(working_histories, type):
        total_experience = timedelta()

        if type == 'All':
            for working_history in working_histories:
                start_date = working_history.startDate
                end_date = working_history.endDate or datetime.now().date()
                experience = end_date - start_date
                total_experience += experience

        if type == 'PravoOhranka':
            for working_history in working_histories:
                if working_history.isPravoOhranka:
                    start_date = working_history.startDate
                    end_date = working_history.endDate or datetime.now().date()
                    experience = end_date - start_date
                    if working_history.HaveCoefficient:
                        experience = experience * 1.5
                    total_experience += experience

        total_years = total_experience.days // 365
        remaining_days = total_experience.days % 365
        total_months = remaining_days // 30
        remaining_days %= 30

        overall_experience = {
            'years': total_years,
            'months': total_months,
            'days': remaining_days
        }

        return overall_experience

    # Load the Word document template
    template_path = 'docx_generator/static/templates/spravka_template.docx'  # Update with the path to your template
    document = Document(template_path)
    tables = document.tables
    p2 = tables[0].rows[0].cells[2].add_paragraph()
    r2 = p2.add_run()
    r2.add_picture(image, width=Inches(1.2))

    # Define a function to replace placeholders in the document
    def replace_placeholder(placeholder, replacement):
        for paragraph1 in tables[0].rows[0].cells[0].paragraphs:
            if placeholder in paragraph1.text:
                for run1 in paragraph1.runs:
                    if placeholder in run1.text:
                        run1.text = run1.text.replace(placeholder, replacement)
                        run1.font.size = Pt(12)  # Adjust the font size if needed

    # Replace placeholders with actual data
    replace_placeholder('placeholder', f"{person.firstName}")
    replace_placeholder('surname', f"{person.surname}")
    replace_placeholder('patronymic', f"{person.patronymic}")
    replace_placeholder('nationality', f"{person.nationality}")
    replace_placeholder('position', person.positionInfo.position.positionTitle)
    replace_placeholder('iin', person.iin)
    replace_placeholder('birth_date', str(formatted_date))
    replace_placeholder('region', birth_info.region)
    replace_placeholder('city', birth_info.city)
    replace_placeholder('rank', rankInfo.militaryRank.rankTitle if rankInfo else '')

    if len(education_data) == 0:
        replace_placeholder('education', "Не имеет")
    else:
        replace_placeholder('education',
                            f"окончил(а) {first_education['educationPlace']} в {date_obj.year} году на специальность {first_education['speciality']}")
    # Create a BytesIO object to save the modified document

    if len(academic_degrees_data) == 0:
        replace_placeholder('academicdegree', "Не имеет")
    else:
        replace_placeholder('academicdegree',
                            f"получил(а) {first_academic_degree['academicDegree']} в {first_academic_degree['academicPlace']} в {date_academ_obj.year} году")
    # Create a BytesIO object to save the modified document

    work_history = WorkingHistory.objects.filter(personId=person).order_by('startDate')
    education_history = Education.objects.filter(personId=person).order_by('educationDateIn')

    pravo_experience = calculate_experience(working_histories=work_history,
                                            type='PravoOhranka')

    if pravo_experience['years'] == 1:
        yearString = 'год'
    elif pravo_experience['years'] in [2, 3, 4, 22, 23, 24, 32, 33, 34, 42, 43, 44, 52, 53, 54]:
        yearString = 'года'
    else:
        yearString = 'лет'

    if pravo_experience['months'] == 1:
        monthString = 'месяц'
    elif pravo_experience['months'] in [2, 3, 4]:
        monthString = 'месяца'
    else:
        monthString = 'месяцев'

    if pravo_experience['days'] == 1:
        dayString = 'день'
    elif pravo_experience['days'] in [2, 3, 4]:
        dayString = 'дня'
    else:
        dayString = 'дней'

    if pravo_experience['years'] == 0 and pravo_experience['months'] == 0:
        pravo_experience_string = str(pravo_experience['days']) + ' ' + dayString
    elif pravo_experience['years'] == 0:
        pravo_experience_string = (str(pravo_experience['months']) + ' ' + monthString + ' '
                                   + str(pravo_experience['days']) + ' ' + dayString)
    else:
        pravo_experience_string = str(pravo_experience['years']) + ' ' + yearString + ' ' + str(
            pravo_experience['months']) + ' ' + monthString + ' ' + str(pravo_experience['days']) + ' ' + dayString

    replace_placeholder('pravoexp', pravo_experience_string)

    # Create a new section in the document after a specific keyword
    keyword = "ДЕЯТЕЛЬНОСТЬ"  # Replace with the keyword you want to use
    for paragraph in document.paragraphs:
        if keyword in paragraph.text:
            section = paragraph._element
            break

    # Create a table with 2 columns for work history and education
    num_columns = 2
    table = document.add_table(rows=1, cols=num_columns)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Define the column widths (adjust these as needed)
    table.columns[0].width = Inches(2)  # Date in and Date out
    table.columns[1].width = Inches(3)  # Place of education and work

    # Add a header row to the table
    table.rows[0].cells[0].text = "Дата"
    table.rows[0].cells[1].text = "Место деятельности и специальность"

    # Iterate through all the cells in the header row
    for cell in table.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER

        # Adjust the paragraph spacing to add padding (adjust the values as needed)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(6)  # Add space before text
        paragraph.paragraph_format.space_after = Pt(6)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                # Iterate through all the cells in the table to add borders

    for entry in education_history:
        date_in = entry.educationDateIn.strftime('%d.%m.%Y')
        date_out = entry.educationDateOut.strftime('%d.%m.%Y')
        place = f"{entry.educationPlace}, Speciality: {entry.speciality}"
        table.add_row().cells[0].text = f"{date_in} - {date_out}"
        table.rows[-1].cells[1].text = place

    # Populate the table with work history and education data
    for entry in work_history:
        date_in = entry.startDate.strftime('%d.%m.%Y')
        date_out = None
        if entry.endDate:
            date_out = entry.endDate.strftime('%d.%m.%Y')
        place = f"{entry.organizationName}, должность: {entry.positionName}"
        if date_out is None:
            table.add_row().cells[0].text = f"{date_in} - по настоящее время"
            table.rows[-1].cells[1].text = place
        else:
            table.add_row().cells[0].text = f"{date_in} - {date_out}"
            table.rows[-1].cells[1].text = place

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(6)  # Add space before text
                paragraph.paragraph_format.space_after = Pt(6)

    doc_stream = BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)

    # Prepare the HTTP response with the modified document
    response = HttpResponse(doc_stream.read(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=work_reference.docx'

    return response


@csrf_exempt
def generate_appointment_decree(request):
    if request.method == 'POST':
        try:
            with transaction.atomic():
                body = request.body.decode('utf-8')
                data = json.loads(body)

                decreeDate = data.get('decreeDate')
                forms = data.get('forms', [])
                bases = [base['base'] for base in data.get('bases', [])]
                # Create a Document object
                template_path = 'docx_generator/static/templates/appointment_template.docx'
                document = Document(template_path)

                # Find the index of the paragraph containing "Председатель"
                keyword_index = -1
                keyword_index_kz = -1
                for i, paragraph in enumerate(document.paragraphs):
                    if "Председатель" in paragraph.text:
                        keyword_index = i
                        break

                for i, paragraph in enumerate(document.paragraphs):
                    for run in paragraph.runs:
                        if "Төраға" in run.text and run.bold:
                            keyword_index_kz = i
                            break

                decree_list_instance = DecreeList.objects.create(
                    decreeType="Назначение",
                    decreeDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                    minioDocName=None
                )

                # If the keyword is found, insert form text after it
                if keyword_index != -1 and keyword_index_kz != -1:
                    last_index = None
                    last_index_kz = None
                    for index, form in enumerate(forms, start=1):
                        person_id = form.get('personId')
                        month_count = form.get('monthCount')
                        appointmentType = form.get('appointmentType')
                        try:
                            personInstance = Person.objects.get(pk=person_id)
                        except Person.DoesNotExist:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': 'Выбранного сотрудника не существует'}, status=400)

                        departmentInstance = personInstance.positionInfo.department
                        positionInstance = personInstance.positionInfo.position
                        positionTitle = positionInstance.positionTitle
                        departmentName = departmentInstance.DepartmentName

                        if AppointmentInfo.objects.filter(personId=personInstance, decreeId__isConfirmed=False,
                                                          decreeId__decreeType="Назначение"):
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'У сотрудника {personInstance.iin} уже существует приказ о назначении который '
                                         f'не согласован'},
                                status=400)

                        try:
                            staffing_table_instance = StaffingTable.objects.get(
                                staffing_table_position=positionInstance,
                                staffing_table_department=departmentInstance)
                        except StaffingTable.DoesNotExist:
                            # If StaffingTable instance doesn't exist, there are no vacancies
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {
                                    'error': f'В базе данных нету штатного расписания с департаментом {departmentInstance.DepartmentName} и должностью {positionInstance.positionTitle}'},
                                status=400)

                        if not staffing_table_instance.vacancy_list.filter(position=positionInstance,
                                                                           department=departmentInstance).first():
                            decree_list_instance.delete()
                            return JsonResponse({
                                'error': f'Нет свободных вакансий для назначения {departmentInstance.DepartmentName} - {positionInstance.positionTitle}'},
                                status=400)

                        if appointmentType == "Впервые принятый":
                            tz = pytz.timezone('Etc/GMT-6')

                            three_months_later = datetime.strptime(decreeDate, '%Y-%m-%d').date() + timedelta(
                                days=int(month_count) * 30 + 1)
                            print(three_months_later)
                            if personInstance.rankInfo is None:

                                appointmentInstance = AppointmentInfo.objects.create(
                                    appointmentDepartment=departmentInstance,
                                    appointmentPosition=positionInstance,
                                    appointmentProbation=int(month_count),
                                    appointmentType=appointmentType,
                                    personId=personInstance,
                                    decreeId=decree_list_instance
                                )

                                task = create_rank_info_after_months.apply_async(
                                    args=(int(month_count), appointmentInstance.id), eta=three_months_later)

                            else:
                                decree_list_instance.delete()
                                return JsonResponse(
                                    {'error': f'Сотрудник {personInstance.iin} уже имеет звание'}, status=400)

                        if appointmentType == "Вновь принятый":
                            AppointmentInfo.objects.create(
                                appointmentDepartment=departmentInstance,
                                appointmentPosition=positionInstance,
                                appointmentProbation=None,
                                appointmentType=appointmentType,
                                personId=personInstance,
                                decreeId=decree_list_instance
                            )
                        soglasnie = ['б', 'в', 'г', 'д', 'ж', 'з', 'й', 'к', 'л', 'м', 'н', 'п', 'р', 'с', 'т', 'ф',
                                     'х',
                                     'ц',
                                     'ч',
                                     'ш', 'щ']
                        glasnie = ['а', 'е', 'ё', 'и', 'о', 'у', 'ы', 'э', 'ю', 'я']

                        changedSurname = personInstance.surname
                        changedFirstName = personInstance.firstName

                        if personInstance.gender.genderName == 'Мужской':
                            if personInstance.firstName[-1] in soglasnie:
                                changedFirstName = personInstance.firstName + 'а'
                            else:
                                changedFirstName = personInstance.firstName

                            if personInstance.surname[-2:] == 'ев' or personInstance.surname[-2:] == 'ов':
                                changedSurname = personInstance.surname + 'а'
                            else:
                                changedSurname = personInstance.surname

                        if personInstance.gender.genderName == 'Женский':
                            if personInstance.firstName[-1] == 'а' and personInstance.firstName[-2] in soglasnie:
                                changedFirstName = personInstance.firstName[:-1]
                                changedFirstName = changedFirstName + 'у'
                            else:
                                changedFirstName = personInstance.firstName

                            if personInstance.surname[-3:] == 'ева' or personInstance.surname[-3:] == 'ова':
                                changedSurname = personInstance.surname[:-1]
                                changedSurname = changedSurname + 'у'
                            else:
                                changedSurname = personInstance.surname

                        personsFIO = changedSurname + ' ' + changedFirstName + ' ' + personInstance.patronymic
                        personsFIOKaz = personInstance.firstName + ' ' + personInstance.patronymic + ' ' + personInstance.surname
                        changedPositionTitle = positionTitle

                        if positionTitle == 'Руководитель департамента':
                            changedPositionTitle = 'Руководителя департамента'
                        if positionTitle == 'Заместитель руководителя департамента':
                            changedPositionTitle = 'Заместителя руководителя департамента'
                        if positionTitle == 'Руководитель управления':
                            changedPositionTitle = 'Руководителя управления'
                        if positionTitle == 'Заместитель руководителя управления':
                            changedPositionTitle = 'Заместителя руководителя управления'
                        if positionTitle == 'Оперуполномоченный по особо важным делам':
                            changedPositionTitle = 'Оперуполномоченного по особо важным делам'
                        if positionTitle == 'Старший оперуполномоченный':
                            changedPositionTitle = 'Старшего оперуполномоченного'
                        if positionTitle == 'Оперуполномоченный':
                            changedPositionTitle = 'Оперуполномоченного'

                        changedDepartmentName = departmentName
                        changedDepartmentNameKz = departmentInstance.DepartmentNameKaz
                        words = departmentName.split()
                        if words[0] == 'Управление':
                            words[0] = 'Управления'
                            changedDepartmentName = ' '.join(words)
                        if departmentName == 'ЦА':
                            changedDepartmentName = 'Управления'
                        if departmentName == 'ЦА':
                            departmentName = 'Управление'

                        if departmentInstance.DepartmentNameKaz == 'Басқарма':
                            changedDepartmentNameKz = 'Басқармасының'
                        else:
                            changedDepartmentNameKz = departmentInstance.DepartmentNameKaz + 'ның'

                        # Bases
                        form_text = None
                        form_text_kz = None

                        if appointmentType == "Впервые принятый":
                            form_text = f"\t{index}. Принять {personsFIO} на правоохранительную службу, назначить на должность {changedPositionTitle.lower()} {changedDepartmentName} _____________________ Агентства Республики Казахстан по финансовому мониторингу, установить испытательный срок {month_count} месяца."
                            form_text_kz = (f"\t{index}. {personsFIOKaz} Қазақстан Республикасы Қаржылық мониторинг "
                                            f"агенттігінің құқық қорғау қызметіне қабылдансын, Агенттіктің "
                                            f"____________________ департаменті "
                                            f"{changedDepartmentNameKz} {positionInstance.positionTitleKaz.lower()} лауазымына "
                                            f"тағайындалсын, {month_count} ай сынақ мерзімі белгіленсін.")
                        if appointmentType == "Вновь принятый":
                            form_text = f"\t{index}. Принять {personsFIO} на правоохранительную службу, назначить на должность {changedPositionTitle.lower()} {changedDepartmentName} _____________________ Агентства Республики Казахстан по финансовому мониторингу."
                            form_text_kz = (f"\t{index}. {personsFIOKaz} Қазақстан Республикасы Қаржылық мониторинг "
                                            f"агенттігінің құқық қорғау қызметіне қабылдансын, Агенттіктің "
                                            f"____________________ департаменті "
                                            f"{changedDepartmentNameKz} {positionInstance.positionTitleKaz.lower()} лауазымына "
                                            f"тағайындалсын.")

                        for i, paragraph in enumerate(document.paragraphs):
                            if "Председатель" in paragraph.text:
                                keyword_index = i
                                break

                        new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(form_text)
                        keyword_index += 1

                        run = new_paragraph.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph.paragraph_format.line_spacing = Pt(16)
                        new_paragraph.paragraph_format.space_after = Pt(0)

                        new_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        last_index = index

                        for i, paragraph in enumerate(document.paragraphs):
                            for run in paragraph.runs:
                                if "Төраға" in run.text and run.bold:
                                    print("helloooo")
                                    keyword_index_kz = i
                                    print(keyword_index_kz)
                                    break

                        new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(form_text_kz)
                        keyword_index_kz += 1

                        run = new_paragraph_kz.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                        new_paragraph_kz.paragraph_format.space_after = Pt(0)

                        new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        new_paragraph.paragraph_format.keep_together = True

                        last_index_kz = index

                    bases_kz = []
                    for base in bases:
                        if base == 'представление':
                            try:
                                predstavlenie = Base.objects.get(baseName="Представление")
                                decree_list_instance.decreeBases.add(predstavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Представление не было найдено в базе данных'},
                                                    status=400)

                            bases_kz.append('ұсыныс')
                        elif base == 'рапорт':
                            try:
                                raport = Base.objects.get(baseName="Рапорт")
                                decree_list_instance.decreeBases.add(raport)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Рапорт не было найден в базе данных'}, status=400)

                            bases_kz.append('баянат')
                        elif base == 'заявление':
                            try:
                                zayavlenie = Base.objects.get(baseName="Заявление")
                                decree_list_instance.decreeBases.add(zayavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Заявление не было найдено в базе данных'}, status=400)

                            bases_kz.append('өтініш')
                        elif base == 'протокол':
                            try:
                                protocol = Base.objects.get(baseName="Протокол")
                                decree_list_instance.decreeBases.add(protocol)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Протокол не было найден в базе данных'}, status=400)
                            bases_kz.append('хаттама')
                    if len(forms) > 1:
                        # Modify bases if needed
                        modified_bases = []
                        modified_bases_kz = []
                        for base in bases:
                            if base == 'представление':
                                modified_bases.append('представления')
                                modified_bases_kz.append('ұсыныстар')
                            elif base == 'рапорт':
                                modified_bases.append('рапорта')
                                modified_bases_kz.append('баянаттар')
                            elif base == 'заявление':
                                modified_bases.append('заявления')
                                modified_bases_kz.append('өтініштер')
                            elif base == 'протокол':
                                modified_bases.append('протокола')
                                modified_bases_kz.append('хаттамалар')
                        bases = modified_bases
                        bases_kz = modified_bases_kz

                    base_text = f"\t{last_index + 1}. Настоящий приказ вступает в силу со дня подписания.\n\tОснование: {', '.join(bases)}."
                    base_text_kz = (f"\t{last_index_kz + 1}. Осы бұйрық қол қойылған күнінен бастап күшіне "
                                    f"енеді.\n\tНегіздеме: {', '.join(bases_kz)}.")

                    # Основание

                    for i, paragraph in enumerate(document.paragraphs):
                        if "Председатель" in paragraph.text:
                            keyword_index = i
                            break

                    new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(base_text)
                    keyword_index += 1

                    run = new_paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph.paragraph_format.line_spacing = Pt(16)
                    new_paragraph.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index].insert_paragraph_before('\n')

                    # Негіздеме

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Төраға" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(base_text_kz)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index_kz].insert_paragraph_before('\n')

                    doc_stream = BytesIO()
                    document.save(doc_stream)
                    doc_stream.seek(0)

                    # Prepare the HTTP response with the modified document
                    response = HttpResponse(doc_stream.read(),
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml'
                                                         '.document')
                    response['Content-Disposition'] = f'attachment; filename=Приказ о назначении.docx'

                    doc_stream.seek(0)
                    document_id = str(uuid4())
                    document_name = f"document_{document_id}.docx"

                    minio_client = Minio(MINIO_ENDPOINT,
                                         access_key=MINIO_ACCESS_KEY,
                                         secret_key=MINIO_SECRET_KEY,
                                         secure=False)

                    minio_client.put_object(MINIO_BUCKET_NAME, document_name, data=doc_stream,
                                            length=len(doc_stream.getvalue()))
                    document_url = f"{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{document_name}"
                    print(document_url)
                    decree_list_instance.minioDocName = document_name
                    decree_list_instance.save()
                    return response

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Упс, что-то пошло не так'}, status=405)


@csrf_exempt
def generate_transfer_decree(request):
    if request.method == 'POST':
        try:
            with transaction.atomic():
                body = request.body.decode('utf-8')
                data = json.loads(body)

                decreeDate = data.get('decreeDate')
                forms = data.get('forms', [])
                bases = [base['base'] for base in data.get('bases', [])]

                template_path = 'docx_generator/static/templates/transfer_template.docx'
                document = Document(template_path)

                # Find the index of the paragraph containing "Председатель"
                keyword_index = -1
                keyword_index_kz = -1
                for i, paragraph in enumerate(document.paragraphs):
                    if "Председатель" in paragraph.text:
                        keyword_index = i
                        break

                for i, paragraph in enumerate(document.paragraphs):
                    for run in paragraph.runs:
                        if "Төраға" in run.text and run.bold:
                            keyword_index_kz = i
                            break

                decree_list_instance = DecreeList.objects.create(
                    decreeType="Перемещение",
                    decreeDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                    minioDocName=None
                )

                if keyword_index != -1 and keyword_index_kz != -1:
                    last_index = None
                    last_index_kz = None
                    vacancy_count = None
                    for index, form in enumerate(forms, start=1):
                        person_id = form.get('personId')
                        newPositionTitle = form.get('newPosition')
                        newDepartmentName = form.get('newDepartment')

                        try:
                            personInstance = Person.objects.get(pk=person_id)
                        except Person.DoesNotExist:
                            decree_list_instance.delete()
                            return JsonResponse({'error': 'Выбранного сотрудника не существует'}, status=400)

                        newDepartmentInstance = Department.objects.get(DepartmentName=newDepartmentName)
                        newPositionInstance = Position.objects.get(positionTitle=newPositionTitle)
                        currentPosition = PositionInfo.objects.get(person=personInstance).position
                        currentDepartment = PositionInfo.objects.get(person=personInstance).department

                        if TransferInfo.objects.filter(personId=personInstance, decreeId__isConfirmed=False,
                                                       decreeId__decreeType="Перемещение"):
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'У сотрудника {personInstance.iin} уже существует приказ о перемещении '
                                         f'который '
                                         f'не согласован'},
                                status=400)

                        if newDepartmentInstance == currentDepartment and newPositionInstance == currentPosition:
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'Сотрудник {personInstance.iin} уже находится на этой позиции'},
                                status=400)

                        try:
                            staffingTableInstance = StaffingTable.objects.get(
                                staffing_table_department=newDepartmentInstance,
                                staffing_table_position=newPositionInstance)
                        except StaffingTable.DoesNotExist:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {
                                    'error': f'В базе данных нету штатного расписания с департаментом {newDepartmentInstance.DepartmentName} и должностью {newPositionInstance.positionTitle}'},
                                status=400)

                        vacancy_count = staffingTableInstance.vacancy_counter
                        print(vacancy_count)
                        if vacancy_count == 0:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {'error': 'На новой должности не существует либо недостаточно актуальной вакансии'},
                                status=400)
                        staffingTableInstance.vacancy_counter = staffingTableInstance.vacancy_counter - 1
                        staffingTableInstance.save()

                        TransferInfo.objects.create(
                            previousDepartment=currentDepartment,
                            previousPosition=currentPosition,
                            newDepartment=newDepartmentInstance,
                            newPosition=newPositionInstance,
                            personId=personInstance,
                            decreeId=decree_list_instance
                        )

                        soglasnie = ['б', 'в', 'г', 'д', 'ж', 'з', 'й', 'к', 'л', 'м', 'н', 'п', 'р', 'с', 'т', 'ф',
                                     'х',
                                     'ц',
                                     'ч',
                                     'ш', 'щ']
                        glasnie = ['а', 'е', 'ё', 'и', 'о', 'у', 'ы', 'э', 'ю', 'я']

                        changedSurname = None
                        changedFirstName = None

                        if personInstance.gender.genderName == 'Мужской':
                            if personInstance.firstName[-1] in soglasnie:
                                changedFirstName = personInstance.firstName + 'а'  # Қасымбаева Қуаныша Ахатұлы
                            else:
                                changedFirstName = personInstance.firstName

                            if personInstance.surname[-2:] == 'ев' or personInstance.surname[-2:] == 'ов':
                                changedSurname = personInstance.surname + 'а'  # Қасымбаева Қуаныша Ахатұлы
                            else:
                                changedSurname = personInstance.surname

                        if personInstance.gender.genderName == 'Женский':
                            if personInstance.firstName[-1] == 'а' and personInstance.firstName[-2] in soglasnie:
                                changedFirstName = personInstance.firstName[:-1]
                                changedFirstName = changedFirstName + 'у'
                            else:
                                changedFirstName = personInstance.firstName

                            if personInstance.surname[-3:] == 'ева' or personInstance.surname[-3:] == 'ова':
                                changedSurname = personInstance.surname[:-1]
                                changedSurname = changedSurname + 'у'
                            else:
                                changedSurname = personInstance.surname

                        personsFIO = changedSurname + ' ' + changedFirstName + ' ' + personInstance.patronymic
                        personsFIOKaz = personInstance.firstName + ' ' + personInstance.patronymic + ' ' + personInstance.surname

                        changedPositionTitle = newPositionInstance.positionTitle
                        changedCurrentPositionTitle = currentPosition.positionTitle

                        if newPositionTitle == 'Руководитель департамента':
                            changedPositionTitle = 'Руководителя департамента'
                        if newPositionTitle == 'Заместитель руководителя департамента':
                            changedPositionTitle = 'Заместителя руководителя департамента'
                        if newPositionTitle == 'Руководитель управления':
                            changedPositionTitle = 'Руководителя управления'
                        if newPositionTitle == 'Заместитель руководителя управления':
                            changedPositionTitle = 'Заместителя руководителя управления'
                        if newPositionTitle == 'Оперуполномоченный по особо важным делам':
                            changedPositionTitle = 'Оперуполномоченного по особо важным делам'
                        if newPositionTitle == 'Старший оперуполномоченный':
                            changedPositionTitle = 'Старшего оперуполномоченного'
                        if newPositionTitle == 'Оперуполномоченный':
                            changedPositionTitle = 'Оперуполномоченного'

                        if currentPosition.positionTitle == 'Руководитель департамента':
                            changedCurrentPositionTitle = 'Руководителя департамента'
                        if currentPosition.positionTitle == 'Заместитель руководителя департамента':
                            changedCurrentPositionTitle = 'Заместителя руководителя департамента'
                        if currentPosition.positionTitle == 'Руководитель управления':
                            changedCurrentPositionTitle = 'Руководителя управления'
                        if currentPosition.positionTitle == 'Заместитель руководителя управления':
                            changedCurrentPositionTitle = 'Заместителя руководителя управления'
                        if currentPosition.positionTitle == 'Оперуполномоченный по особо важным делам':
                            changedCurrentPositionTitle = 'Оперуполномоченного по особо важным делам'
                        if currentPosition.positionTitle == 'Старший оперуполномоченный':
                            changedCurrentPositionTitle = 'Старшего оперуполномоченного'
                        if currentPosition.positionTitle == 'Оперуполномоченный':
                            changedCurrentPositionTitle = 'Оперуполномоченного'

                        changedDepartmentName = newDepartmentInstance.DepartmentName
                        changedCurrentDepartmentName = currentDepartment.DepartmentName
                        changedCurrentDepartmentNameKaz = currentDepartment.DepartmentNameKaz
                        changedNewDepartmentNameKaz = newDepartmentInstance.DepartmentNameKaz

                        words = newDepartmentName.split()
                        if words[0] == 'Управление':
                            words[0] = 'Управления'
                            changedDepartmentName = ' '.join(words)
                        if newDepartmentName == 'ЦА':
                            changedDepartmentName = 'Управления'
                        if newDepartmentName == 'ЦА':
                            departmentName = 'Управление'

                        words = currentDepartment.DepartmentName.split()
                        if words[0] == 'Управление':
                            words[0] = 'Управления'
                            changedCurrentDepartmentName = ' '.join(words)
                        if currentDepartment.DepartmentName == 'ЦА':
                            changedCurrentDepartmentName = 'Управления'
                        if currentDepartment.DepartmentName == 'ЦА':
                            currentDepartment.DepartmentName = 'Управление'

                        if currentDepartment.DepartmentNameKaz == 'Басқарма':
                            changedCurrentDepartmentNameKaz = 'Басқармасының'
                        else:
                            changedCurrentDepartmentNameKaz = currentDepartment.DepartmentNameKaz + 'ның'

                        if newDepartmentInstance.DepartmentNameKaz == 'Басқарма':
                            changedNewDepartmentNameKaz = 'Басқармасының'
                        else:
                            changedNewDepartmentNameKaz = newDepartmentInstance.DepartmentNameKaz + 'ның'

                        form_text = None
                        form_text_kz = None

                        form_text = f"\t{index}. Назначить {personsFIO} на должность {changedPositionTitle.lower()} {changedDepartmentName} _______________________ департамента Агентства Республики Казахстан по финансовому мониторингу (далее - Агентство), освободив от занимаемой должности {changedCurrentPositionTitle.lower()} {changedCurrentDepartmentName} ______________________ департамента Агентства."
                        form_text_kz = f"\t{index}. {personsFIOKaz} атқарып отырған Қазақстан Республикасы Қаржылық мониторинг агенттігінің (бұдан әрі - Агенттік) _______________________ департаменті {changedCurrentDepartmentNameKaz} {currentPosition.positionTitleKaz.lower()} лауазымынан босатылып, Агенттіктің _____________________ департаменті {changedNewDepartmentNameKaz} {newPositionInstance.positionTitleKaz.lower()} лауазымына тағайындалсын."

                        for i, paragraph in enumerate(document.paragraphs):
                            if "Председатель" in paragraph.text:
                                keyword_index = i
                                break

                        new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(form_text)
                        keyword_index += 1

                        run = new_paragraph.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph.paragraph_format.line_spacing = Pt(16)
                        new_paragraph.paragraph_format.space_after = Pt(0)

                        new_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        last_index = index

                        for i, paragraph in enumerate(document.paragraphs):
                            for run in paragraph.runs:
                                if "Төраға" in run.text and run.bold:
                                    keyword_index_kz = i
                                    break

                        new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(form_text_kz)
                        keyword_index_kz += 1

                        run = new_paragraph_kz.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                        new_paragraph_kz.paragraph_format.space_after = Pt(0)

                        new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        new_paragraph.paragraph_format.keep_together = True

                        last_index_kz = index

                        vacancy_count -= 1

                    bases_kz = []
                    for base in bases:
                        if base == 'представление':
                            try:
                                predstavlenie = Base.objects.get(baseName="Представление")
                                decree_list_instance.decreeBases.add(predstavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Представление не было найдено в базе данных'},
                                                    status=400)

                            bases_kz.append('ұсыныс')
                        elif base == 'рапорт':
                            try:
                                raport = Base.objects.get(baseName="Рапорт")
                                decree_list_instance.decreeBases.add(raport)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Рапорт не было найден в базе данных'}, status=400)

                            bases_kz.append('баянат')
                        elif base == 'заявление':
                            try:
                                zayavlenie = Base.objects.get(baseName="Заявление")
                                decree_list_instance.decreeBases.add(zayavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Заявление не было найдено в базе данных'}, status=400)

                            bases_kz.append('өтініш')
                        elif base == 'протокол':
                            try:
                                protocol = Base.objects.get(baseName="Протокол")
                                decree_list_instance.decreeBases.add(protocol)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Протокол не было найден в базе данных'}, status=400)
                            bases_kz.append('хаттама')
                    if len(forms) > 1:
                        # Modify bases if needed
                        modified_bases = []
                        modified_bases_kz = []
                        for base in bases:
                            if base == 'представление':
                                modified_bases.append('представления')
                                modified_bases_kz.append('ұсыныстар')
                            elif base == 'рапорт':
                                modified_bases.append('рапорта')
                                modified_bases_kz.append('баянаттар')
                            elif base == 'заявление':
                                modified_bases.append('заявления')
                                modified_bases_kz.append('өтініштер')
                            elif base == 'протокол':
                                modified_bases.append('протокола')
                                modified_bases_kz.append('хаттамалар')
                        bases = modified_bases
                        bases_kz = modified_bases_kz

                    base_text = f"\t{last_index + 1}. Настоящий приказ вступает в силу со дня подписания.\n\tОснование: {', '.join(bases)}."
                    base_text_kz = (f"\t{last_index_kz + 1}. Осы бұйрық қол қойылған күнінен бастап күшіне "
                                    f"енеді.\n\tНегіздеме: {', '.join(bases_kz)}.")

                    # Основание

                    for i, paragraph in enumerate(document.paragraphs):
                        if "Председатель" in paragraph.text:
                            keyword_index = i
                            break

                    new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(base_text)
                    keyword_index += 1

                    run = new_paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph.paragraph_format.line_spacing = Pt(16)
                    new_paragraph.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index].insert_paragraph_before('\n')

                    # Негіздеме

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Төраға" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(base_text_kz)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index_kz].insert_paragraph_before('\n')

                    doc_stream = BytesIO()
                    document.save(doc_stream)
                    doc_stream.seek(0)

                    # Prepare the HTTP response with the modified document
                    response = HttpResponse(doc_stream.read(),
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml'
                                                         '.document')
                    response['Content-Disposition'] = f'attachment; filename=Приказ о перемещении.docx'

                    doc_stream.seek(0)
                    document_id = str(uuid4())
                    document_name = f"document_{document_id}.docx"

                    minio_client = Minio(MINIO_ENDPOINT,
                                         access_key=MINIO_ACCESS_KEY,
                                         secret_key=MINIO_SECRET_KEY,
                                         secure=False)

                    minio_client.put_object(MINIO_BUCKET_NAME, document_name, data=doc_stream,
                                            length=len(doc_stream.getvalue()))
                    document_url = f"{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{document_name}"
                    print(document_url)
                    decree_list_instance.minioDocName = document_name
                    decree_list_instance.save()
                    return response

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Упс, что-то пошло не так'}, status=405)


@csrf_exempt
def generate_rankup_decree(request):
    if request.method == 'POST':
        try:
            with transaction.atomic():
                body = request.body.decode('utf-8')
                data = json.loads(body)

                decreeDate = data.get('decreeDate')

                forms = data.get('forms', [])
                bases = [base['base'] for base in data.get('bases', [])]

                template_path = 'docx_generator/static/templates/rankup_template.docx'
                document = Document(template_path)

                keyword_index = -1
                keyword_index_kz = -1

                for i, paragraph in enumerate(document.paragraphs):
                    if "Председатель" in paragraph.text:
                        keyword_index = i
                        break

                for i, paragraph in enumerate(document.paragraphs):
                    for run in paragraph.runs:
                        if "Төраға" in run.text and run.bold:
                            keyword_index_kz = i
                            break

                decree_list_instance = DecreeList.objects.create(
                    decreeType="Присвоение звания",
                    decreeDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                    minioDocName=None
                )

                if keyword_index != -1 and keyword_index_kz != -1:
                    last_index = None
                    last_index_kz = None
                    for index, form in enumerate(forms, start=1):
                        person_id = form.get('personId')
                        newRankTitle = form.get('newRank')
                        rankUpDate = form.get('rankUpDate')
                        receivedType = form.get('receivedType')

                        try:
                            personInstance = Person.objects.get(pk=person_id)
                        except Person.DoesNotExist:
                            decree_list_instance.delete()
                            return JsonResponse({'error': 'Выбранного сотрудника не существует'}, status=400)

                        currentPosition = PositionInfo.objects.get(person=personInstance).position
                        currentDepartment = PositionInfo.objects.get(person=personInstance).department

                        personsRankInfo = RankInfo.objects.get(person=personInstance)
                        personsPositionInfo = PositionInfo.objects.get(person=personInstance)

                        try:
                            newRankInstance = MilitaryRank.objects.get(rankTitle=newRankTitle)
                        except json.JSONDecodeError:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': 'Неправильное звание'}, status=400)

                        if RankUpInfo.objects.filter(personId=personInstance, decreeId__isConfirmed=False,
                                                     decreeId__decreeType="Присвоение звания"):
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'У сотрудника {personInstance.iin} уже существует приказ о присвоении звания '
                                         f'который '
                                         f'не согласован'},
                                status=400)

                        newRankTitle = newRankInstance.rankTitle.lower()

                        if personsRankInfo.militaryRank.order > newRankInstance.order:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': 'Новое звание должно быть выше нынешного звания'}, status=400)

                        if receivedType != 'Внеочередное' and personsRankInfo.militaryRank.order + 1 != newRankInstance.order:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {'error': 'Новое звание должно быть следующим званием а не выше чем на 2 звания'},
                                status=400)

                        if receivedType == 'На одну ступень выше специального звания' and personsPositionInfo.position.maxRank.order < newRankInstance.order:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': 'Новое звание превышает максимальное звание должности'},
                                                status=400)

                        changedRankTitleKaz = newRankTitle
                        if newRankTitle == 'старший лейтенант':
                            changedRankTitleKaz = 'аға лейтенант'

                        if receivedType == 'Досрочное':
                            time_difference = datetime.strptime(rankUpDate,
                                                                "%Y-%m-%d").date() - personsRankInfo.receivedDate
                            half_promotion_days = personsRankInfo.militaryRank.nextPromotionDateInDays / 2
                            if time_difference >= timedelta(days=half_promotion_days):
                                RankUpInfo.objects.create(
                                    previousRank=personsRankInfo.militaryRank,
                                    newRank=newRankInstance,
                                    receivedType=receivedType,
                                    receivedDate=rankUpDate,
                                    personId=personInstance,
                                    decreeId=decree_list_instance
                                )
                            else:
                                transaction.set_rollback(True)
                                return JsonResponse(
                                    {'error': 'Ошибка досрочного повышения: Дата повышения не равно или не '
                                              'превышает половины даты'
                                              ' последующего повышения'}, status=400)

                        if receivedType == 'Внеочередное':
                            if personsRankInfo.militaryRank.rankTitle != 'Подполковник' or personsRankInfo.militaryRank.rankTitle != 'Полковник':
                                rankUpDate_dateformat = datetime.strptime(rankUpDate, "%Y-%m-%d").date()
                                # Checking if rankUpDate is crossed over
                                if personsRankInfo.nextPromotionDate <= rankUpDate_dateformat:
                                    # Taking promotion days count from current and next rank to get how many days
                                    # needed to RankUp Внеочередное
                                    current_promotion_days = personsRankInfo.militaryRank.nextPromotionDateInDays
                                    next_one_step_rank = MilitaryRank.objects.get(
                                        order=personsRankInfo.militaryRank.order + 1)
                                    one_step_promotion_days = next_one_step_rank.nextPromotionDateInDays
                                    oneStepRankUpDate = personsRankInfo.receivedDate + timedelta(
                                        days=current_promotion_days + one_step_promotion_days)
                                    # Checking if given rank is greater than one step of next rank
                                    if next_one_step_rank.order + 1 == newRankInstance.order:
                                        # Checking if rankUpDate with 2 ranks is allowed with given date in request
                                        if oneStepRankUpDate <= rankUpDate_dateformat:

                                            RankUpInfo.objects.create(
                                                previousRank=personsRankInfo.militaryRank,
                                                newRank=newRankInstance,
                                                receivedDate=rankUpDate,
                                                receivedType=receivedType,
                                                personId=personInstance,
                                                decreeId=decree_list_instance
                                            )
                                        else:
                                            transaction.set_rollback(True)
                                            return JsonResponse(
                                                {
                                                    'error': f'Дата внеочередного повышения должна соответствовать '
                                                             f'требованиям:'
                                                             f'предпологаемая дата повышения {oneStepRankUpDate}'},
                                                status=400)
                                    else:
                                        transaction.set_rollback(True)
                                        return JsonResponse(
                                            {'error': 'Новое звание должно быть через одну ступень после '
                                                      'нынешного звания'},
                                            status=400)
                                else:
                                    transaction.set_rollback(True)
                                    return JsonResponse(
                                        {'error': 'Дата повышения в приказе должна быть после даты '
                                                  'следующего'
                                                  'повышения'},
                                        status=400)
                            else:
                                transaction.set_rollback(True)
                                return JsonResponse({'error': 'Нету доступа получения внеочередного повышения'},
                                                    status=400)

                        if receivedType == 'На одну ступень выше специального звания':
                            if personsRankInfo.militaryRank.order + 1 == newRankInstance.order:
                                rankUpDate_dateformat = datetime.strptime(rankUpDate, "%Y-%m-%d").date()
                                if personsRankInfo.nextPromotionDate <= rankUpDate_dateformat:
                                    RankUpInfo.objects.create(
                                        previousRank=personsRankInfo.militaryRank,
                                        newRank=newRankInstance,
                                        receivedType=receivedType,
                                        receivedDate=rankUpDate,
                                        personId=personInstance,
                                        decreeId=decree_list_instance
                                    )
                                else:
                                    transaction.set_rollback(True)
                                    return JsonResponse(
                                        {'error': 'Данная дата не подходит с датой следующего повышения'},
                                        status=400)
                            else:
                                transaction.set_rollback(True)
                                return JsonResponse(
                                    {'error': 'Новое звание должно быть следующим после нынешного звания'},
                                    status=400)

                        soglasnie = ['б', 'в', 'г', 'д', 'ж', 'з', 'й', 'к', 'л', 'м', 'н', 'п', 'р', 'с', 'т', 'ф',
                                     'х', 'ц',
                                     'ч',
                                     'ш', 'щ']
                        glasnie = ['а', 'е', 'ё', 'и', 'о', 'у', 'ы', 'э', 'ю', 'я']

                        changedSurname = personInstance.surname
                        changedFirstName = personInstance.firstName
                        changedSurnameKaz = personInstance.surname
                        changedPatronymic = personInstance.patronymic

                        if personInstance.gender.genderName == 'Мужской':
                            if personInstance.firstName[-1] in soglasnie:
                                changedFirstName = personInstance.firstName + 'у'
                            else:
                                changedFirstName = personInstance.firstName

                            if personInstance.surname[-2:] == 'ев' or personInstance.surname[-2:] == 'ов':
                                changedSurname = personInstance.surname + 'у'
                                changedSurnameKaz = personInstance.surname + 'қа'
                            else:
                                changedSurname = personInstance.surname
                                changedSurnameKaz = personInstance.surname

                        if personInstance.gender.genderName == 'Женский':
                            if personInstance.firstName[-1] == 'а':
                                changedFirstName = personInstance.firstName[:-1]
                                changedFirstName = changedFirstName + 'е'
                            else:
                                changedFirstName = personInstance.firstName

                            if personInstance.surname[-3:] == 'ева' or personInstance.surname[-3:] == 'ова':
                                changedSurname = personInstance.surname[:-1]
                                changedSurname = changedSurname + 'ой'
                                changedSurnameKaz = personInstance.surname + 'ға'
                            else:
                                changedSurname = personInstance.surname

                            if personInstance.patronymic[-4:] == 'евна' or personInstance.patronymic[-4:] == 'овна':
                                changedPatronymic = personInstance.patronymic[:-1]
                                changedPatronymic = changedPatronymic + 'е'
                            else:
                                changedPatronymic = personInstance.patronymic

                        personsFIO = changedSurname + ' ' + changedFirstName + ' ' + changedPatronymic
                        personsFIOKaz = personInstance.firstName + ' ' + personInstance.patronymic + ' ' + changedSurnameKaz

                        changedCurrentDepartmentName = currentDepartment.DepartmentName
                        changedCurrentDepartmentNameKaz = currentDepartment.DepartmentNameKaz
                        changedCurrentPositionKaz = currentPosition.positionTitleKaz
                        changedCurrentPosition = currentPosition.positionTitle

                        words = currentDepartment.DepartmentName.split()
                        if words[0] == 'Управление':
                            words[0] = 'Управления'
                            changedCurrentDepartmentName = ' '.join(words)

                        if currentDepartment.DepartmentName == 'ЦА':
                            changedCurrentDepartmentName = 'Управления'

                        if currentDepartment.DepartmentName == 'ЦА':
                            currentDepartment.DepartmentName = 'Управление'

                        if currentDepartment.DepartmentNameKaz == 'Басқарма':
                            changedCurrentDepartmentNameKaz = 'Басқармасының'
                        else:
                            changedCurrentDepartmentNameKaz = currentDepartment.DepartmentNameKaz + 'ның'

                        if currentPosition.positionTitle == 'Руководитель департамента':
                            changedCurrentPosition = 'Руководителю департамента'

                        if currentPosition.positionTitle == 'Заместитель руководителя департамента':
                            changedCurrentPosition = 'Заместителю руководителя департамента'

                        if currentPosition.positionTitle == 'Руководитель управления':
                            changedCurrentPosition = 'Руководителю управления'

                        if currentPosition.positionTitle == 'Заместитель руководителя управления':
                            changedCurrentPosition = 'Заместителю руководителя управления'

                        if currentPosition.positionTitle == 'Оперуполномоченный по особо важным делам':
                            changedCurrentPosition = 'Оперуполномоченному по особо важным делам'

                        if currentPosition.positionTitle == 'Старший оперуполномоченный':
                            changedCurrentPosition = 'Старшему оперуполномоченному'

                        if currentPosition.positionTitle == 'Оперуполномоченный':
                            changedCurrentPosition = 'Оперуполномоченному'

                        if currentPosition.positionTitleKaz == 'Аға жедел уәкіл':
                            changedCurrentPositionKaz = 'Аға жедел уәкілі'
                        if currentPosition.positionTitleKaz == 'Жедел уәкіл':
                            changedCurrentPositionKaz = 'Жедел уәкілі'
                        if currentPosition.positionTitleKaz == 'Аса маңызды істер жөніндегі жедел уәкіл':
                            changedCurrentPositionKaz = 'Аса маңызды істер жөніндегі жедел уәкілі'

                        changedCurrentPosition = changedCurrentPosition.lower()
                        changedCurrentPositionKaz = changedCurrentPositionKaz.lower()

                        year, month, day = map(int, rankUpDate.split('-'))
                        day = int(day)
                        monthString = None
                        monthStringKaz = None

                        if month == 1:
                            monthString = 'января'
                            monthStringKaz = 'қантардан'
                        if month == 2:
                            monthString = 'февраля'
                            monthStringKaz = 'ақпаннан'
                        if month == 3:
                            monthString = 'марта'
                            monthStringKaz = 'наурыздан'
                        if month == 4:
                            monthString = 'апреля'
                            monthStringKaz = 'сәуірден'
                        if month == 5:
                            monthString = 'мая'
                            monthStringKaz = 'мамырдан'
                        if month == 6:
                            monthString = 'июня'
                            monthStringKaz = 'маусымнан'
                        if month == 7:
                            monthString = 'июля'
                            monthStringKaz = 'шілдеден'
                        if month == 8:
                            monthString = 'августа'
                            monthStringKaz = 'тамыздан'
                        if month == 9:
                            monthString = 'сентября'
                            monthStringKaz = 'қыркүйектен'
                        if month == 10:
                            monthString = 'октября'
                            monthStringKaz = 'қазаннан'
                        if month == 11:
                            monthString = 'ноября'
                            monthStringKaz = 'қарашадан'
                        if month == 12:
                            monthString = 'декабря'
                            monthStringKaz = 'желтоқсаннан'

                        changedDateKaz = str(year) + ' жылғы ' + str(day) + ' ' + monthStringKaz
                        changedDate = str(day) + ' ' + monthString + ' ' + str(year) + ' ' + 'года'

                        form_text = None
                        form_text_kz = None

                        form_text = f"\t{index}. Присвоить очередное специальное звание {newRankTitle} службы экономических расследований {changedCurrentPosition} {changedCurrentDepartmentName} ________________________ департамента Агентства Республики Казахстан по финансовому мониторингу {personsFIO} с {changedDate}."
                        form_text_kz = f"\t{index}. Кезекті экономикалық тергеп-тексеру қызметінің {changedRankTitleKaz} арнаулы атағы Қазақстан Республикасы Қаржылық мониторинг агенттігінің ________________________ департаменті {changedCurrentDepartmentNameKaz} {changedCurrentPositionKaz} {personsFIOKaz} {changedDateKaz} бастап берілсін."

                        for i, paragraph in enumerate(document.paragraphs):
                            if "Председатель" in paragraph.text:
                                keyword_index = i
                                break

                        new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(form_text)
                        keyword_index += 1

                        run = new_paragraph.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph.paragraph_format.line_spacing = Pt(16)
                        new_paragraph.paragraph_format.space_after = Pt(0)

                        new_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        last_index = index

                        for i, paragraph in enumerate(document.paragraphs):
                            for run in paragraph.runs:
                                if "Төраға" in run.text and run.bold:
                                    keyword_index_kz = i
                                    break

                        new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(form_text_kz)
                        keyword_index_kz += 1

                        run = new_paragraph_kz.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                        new_paragraph_kz.paragraph_format.space_after = Pt(0)

                        new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        new_paragraph.paragraph_format.keep_together = True

                        last_index_kz = index

                    bases_kz = []
                    for base in bases:
                        if base == 'представление':
                            try:
                                predstavlenie = Base.objects.get(baseName="Представление")
                                decree_list_instance.decreeBases.add(predstavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Представление не было найдено в базе данных'},
                                                    status=400)

                            bases_kz.append('ұсыныс')
                        elif base == 'рапорт':
                            try:
                                raport = Base.objects.get(baseName="Рапорт")
                                decree_list_instance.decreeBases.add(raport)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Рапорт не было найден в базе данных'}, status=400)

                            bases_kz.append('баянат')
                        elif base == 'заявление':
                            try:
                                zayavlenie = Base.objects.get(baseName="Заявление")
                                decree_list_instance.decreeBases.add(zayavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Заявление не было найдено в базе данных'}, status=400)

                            bases_kz.append('өтініш')
                        elif base == 'протокол':
                            try:
                                protocol = Base.objects.get(baseName="Протокол")
                                decree_list_instance.decreeBases.add(protocol)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Протокол не было найден в базе данных'}, status=400)
                            bases_kz.append('хаттама')
                    if len(forms) > 1:
                        # Modify bases if needed
                        modified_bases = []
                        modified_bases_kz = []
                        for base in bases:
                            if base == 'представление':
                                modified_bases.append('представления')
                                modified_bases_kz.append('ұсыныстар')
                            elif base == 'рапорт':
                                modified_bases.append('рапорта')
                                modified_bases_kz.append('баянаттар')
                            elif base == 'заявление':
                                modified_bases.append('заявления')
                                modified_bases_kz.append('өтініштер')
                            elif base == 'протокол':
                                modified_bases.append('протокола')
                                modified_bases_kz.append('хаттамалар')
                        bases = modified_bases
                        bases_kz = modified_bases_kz

                    base_text = f"\t{last_index + 1}. Настоящий приказ вступает в силу со дня подписания.\n\tОснование: {', '.join(bases)}."
                    base_text_kz = (f"\t{last_index_kz + 1}. Осы бұйрық қол қойылған күнінен бастап күшіне "
                                    f"енеді.\n\tНегіздеме: {', '.join(bases_kz)}.")

                    # Основание

                    for i, paragraph in enumerate(document.paragraphs):
                        if "Председатель" in paragraph.text:
                            keyword_index = i
                            break

                    new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(base_text)
                    keyword_index += 1

                    run = new_paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph.paragraph_format.line_spacing = Pt(16)
                    new_paragraph.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index].insert_paragraph_before('\n')

                    # Негіздеме

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Төраға" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(base_text_kz)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index_kz].insert_paragraph_before('\n')

                    doc_stream = BytesIO()
                    document.save(doc_stream)
                    doc_stream.seek(0)

                    # Prepare the HTTP response with the modified document
                    response = HttpResponse(doc_stream.read(),
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml'
                                                         '.document')
                    response['Content-Disposition'] = f'attachment; filename=Приказ о присвоении звания.docx'

                    doc_stream.seek(0)
                    document_id = str(uuid4())
                    document_name = f"document_{document_id}.docx"

                    minio_client = Minio(MINIO_ENDPOINT,
                                         access_key=MINIO_ACCESS_KEY,
                                         secret_key=MINIO_SECRET_KEY,
                                         secure=False)

                    minio_client.put_object(MINIO_BUCKET_NAME, document_name, data=doc_stream,
                                            length=len(doc_stream.getvalue()))
                    document_url = f"{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{document_name}"
                    print(document_url)
                    decree_list_instance.minioDocName = document_name
                    decree_list_instance.save()
                    return response

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Упс, что-то пошло не так'}, status=405)


@csrf_exempt
def generate_firing_decree(request):
    if request.method == 'POST':
        try:
            with transaction.atomic():
                body = request.body.decode('utf-8')
                data = json.loads(body)

                decreeDate = data.get('decreeDate')

                forms = data.get('forms', [])
                bases = [base['base'] for base in data.get('bases', [])]

                template_path = 'docx_generator/static/templates/firing_template.docx'
                document = Document(template_path)

                keyword_index = -1
                keyword_index_kz = -1

                for i, paragraph in enumerate(document.paragraphs):
                    if "Председатель" in paragraph.text:
                        keyword_index = i
                        break

                for i, paragraph in enumerate(document.paragraphs):
                    for run in paragraph.runs:
                        if "Төраға" in run.text and run.bold:
                            keyword_index_kz = i
                            break

                decree_list_instance = DecreeList.objects.create(
                    decreeType="Увольнение",
                    decreeDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                    minioDocName=None
                )

                if keyword_index != -1 and keyword_index_kz != -1:
                    last_index = None
                    last_index_kz = None
                    for index, form in enumerate(forms, start=1):
                        person_id = form.get('personId')

                        try:
                            personInstance = Person.objects.get(pk=person_id)
                        except Person.DoesNotExist:
                            decree_list_instance.delete()
                            return JsonResponse({'error': 'Выбранного сотрудника не существует'}, status=400)

                        personsPositionInfo = PositionInfo.objects.get(person=personInstance)

                        positionTitle = personsPositionInfo.position.positionTitle
                        departmentName = personsPositionInfo.department.DepartmentName

                        date_object = datetime.strptime(decreeDate, "%Y-%m-%d")

                        basicVacancyDays = 0
                        expVacancyDays = 0
                        basicVacation = None
                        expVacation = None

                        try:
                            basicVacation = Vacation.objects.get(personId=personInstance, daysType="Обычные",
                                                                 year=date_object.year)
                        except Vacation.DoesNotExist:
                            pass

                        if basicVacation:
                            basicVacancyDays = basicVacation.daysCount
                        try:
                            expVacation = Vacation.objects.get(personId=personInstance, daysType="Стажные",
                                                               year=date_object.year)
                        except Vacation.DoesNotExist:
                            pass

                        if expVacation:
                            expVacancyDays = expVacation.daysCount

                        dayCount = basicVacancyDays + expVacancyDays

                        if FiringInfo.objects.filter(personId=personInstance, decreeId__isConfirmed=False,
                                                     decreeId__decreeType="Увольнение"):
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'У сотрудника {personInstance.iin} уже существует приказ об увольнении '
                                         f'который '
                                         f'не согласован'},
                                status=400)

                        if personInstance.isFired:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {'error': f'Сотрудник {personInstance.iin} уже уволен'},
                                status=400)

                        FiringInfo.objects.create(
                            firingDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                            personId=personInstance,
                            decreeId=decree_list_instance
                        )

                        soglasnie = ['б', 'в', 'г', 'д', 'ж', 'з', 'й', 'к', 'л', 'м', 'н', 'п', 'р', 'с', 'т', 'ф',
                                     'х', 'ц',
                                     'ч',
                                     'ш', 'щ']
                        glasnie = ['а', 'е', 'ё', 'и', 'о', 'у', 'ы', 'э', 'ю', 'я']

                        # Kassymbayeva Kuanysh Akhatuly
                        changedSurname = personInstance.surname
                        changedFirstName = personInstance.firstName
                        changedPatronymic = personInstance.patronymic

                        # Kassymbayevu Kuanyshu Akhatuly
                        changedSurname2 = personInstance.surname
                        changedFirstName2 = personInstance.firstName
                        changedPatronymic2 = personInstance.patronymic

                        if personInstance.gender.genderName == 'Мужской':
                            if personInstance.firstName[-1] in soglasnie:
                                changedFirstName = personInstance.firstName + 'а'
                                changedFirstName2 = personInstance.firstName + 'у'

                            if personInstance.surname[-2:] == 'ев' or personInstance.surname[-2:] == 'ов':
                                changedSurname = personInstance.surname + 'а'
                                changedSurname2 = personInstance.surname + 'у'

                            if personInstance.patronymic[-3:] == 'вич':
                                changedPatronymic = personInstance.patronymic + 'а'
                                changedPatronymic2 = personInstance.patronymic + 'у'

                        if personInstance.gender.genderName == 'Женский':
                            if personInstance.firstName[-1] == 'а' and personInstance.firstName[-2] in soglasnie:
                                changedFirstName = personInstance.firstName[:-1]
                                changedFirstName2 = personInstance.firstName[:-1]

                                changedFirstName = changedFirstName + 'у'
                                changedFirstName2 = changedFirstName2 + 'е'

                            if personInstance.surname[-3:] == 'ева' or personInstance.surname[-3:] == 'ова':
                                changedSurname = personInstance.surname[:-1]
                                changedSurname2 = personInstance.surname[:-1]

                                changedSurname = changedSurname + 'у'
                                changedSurname2 = changedSurname2 + 'е'

                            if personInstance.patronymic[-4:] == 'овна' or personInstance.patronymic[-4:] == 'евна':
                                changedPatronymic = personInstance.patronymic[:-1]
                                changedPatronymic2 = personInstance.patronymic[:-1]

                                changedPatronymic = changedPatronymic + 'у'
                                changedPatronymic2 = changedPatronymic2 + 'е'

                        personsFIO = changedSurname + ' ' + changedFirstName + ' ' + changedPatronymic
                        personsFIOKaz = personInstance.firstName + ' ' + personInstance.patronymic + ' ' + personInstance.surname

                        personsFIO2 = changedSurname2 + ' ' + changedFirstName2 + ' ' + changedPatronymic2

                        changedPositionTitle = positionTitle
                        if positionTitle == 'Руководитель департамента':
                            changedPositionTitle = 'Руководителя департамента'
                        if positionTitle == 'Заместитель руководителя департамента':
                            changedPositionTitle = 'Заместителя руководителя департамента'
                        if positionTitle == 'Руководитель управления':
                            changedPositionTitle = 'Руководителя управления'
                        if positionTitle == 'Заместитель руководителя управления':
                            changedPositionTitle = 'Заместителя руководителя управления'
                        if positionTitle == 'Оперуполномоченный по особо важным делам':
                            changedPositionTitle = 'Оперуполномоченного по особо важным делам'
                        if positionTitle == 'Старший оперуполномоченный':
                            changedPositionTitle = 'Старшего оперуполномоченного'
                        if positionTitle == 'Оперуполномоченный':
                            changedPositionTitle = 'Оперуполномоченного'

                        changedDepartmentName = departmentName
                        changedDepartmentNameKaz = personsPositionInfo.department.DepartmentNameKaz
                        words = departmentName.split()
                        wordsKaz = changedDepartmentNameKaz.split()
                        if words[0] == 'Управление':
                            words[0] = 'управления'
                            changedDepartmentName = ' '.join(words)
                        if departmentName == 'ЦА':
                            changedDepartmentName = 'управления'
                        if departmentName == 'ЦА':
                            departmentName = 'управление'

                        changedDepartmentName2 = departmentName
                        words = departmentName.split()
                        if words[0] == 'Управление':
                            words[0] = 'Управлению'
                            changedDepartmentName2 = ' '.join(words)
                        if departmentName == 'ЦА':
                            changedDepartmentName2 = 'Управлению'
                        if departmentName == 'ЦА':
                            departmentName = 'Управление'

                        if wordsKaz[-1] == 'басқармасы':
                            wordsKaz[-1] = wordsKaz[-1] + 'ның'
                            changedDepartmentNameKaz = ' '.join(wordsKaz)
                        if wordsKaz[0] == 'Басқарма':
                            wordsKaz[0] = 'басқармасы' + 'ның'
                            changedDepartmentNameKaz = ' '.join(wordsKaz)

                        form_text = None
                        form_text_kz = None

                        if dayCount == 0:
                            form_text = f"\t{index}. Освободить {personsFIO} от занимаемой должности {changedPositionTitle.lower()} {changedDepartmentName} ___________________ департамента Агентства, уволить и исключить из кадров правоохранительного органа – службы экономических расследований органов по финансовому мониторингу по собственному желанию."
                            form_text_kz = f"\t{index}. {personsFIOKaz} атқарып отырған Агенттіктің ___________________ департаменті {changedDepartmentNameKaz} {personsPositionInfo.position.positionTitleKaz.lower()} лауазымынан босатылсын, құқық қорғау органы – қаржылық мониторинг органдарының экономикалық тергеп-тексеру қызметінен және кадрынан өз қалауы бойынша шығарылсын."
                        if dayCount > 0:
                            form_text = f"\t{index}. 1.	Освободить {personsFIO} от занимаемой должности {changedPositionTitle.lower()} {changedDepartmentName} ___________________ департамента Агентства, уволить и исключить из кадров правоохранительного органа – службы экономических расследований органов по финансовому мониторингу по собственному желанию. \n\t{changedDepartmentName2} ______________________ департамента Агентства произвести {personsFIO2} выплату за неиспользованные {dayCount} дней отпуска за {date_object.year} год."
                            form_text_kz = f"\t{index}. {personsFIOKaz} атқарып отырған Агенттіктің ___________________ департаменті {changedDepartmentNameKaz} {personsPositionInfo.position.positionTitleKaz.lower()} лауазымынан босатылсын, құқық қорғау органы – қаржылық мониторинг органдарының экономикалық тергеп-тексеру қызметінен және кадрынан өз қалауы бойынша шығарылсын. \n\tАгенттіктің ______________________ департаментінің {personsPositionInfo.department.DepartmentNameKaz} {personsFIOKaz} {date_object.year} жылғы пайдаланбаған {dayCount} күн еңбек демалысына төлем жүргізсін."

                        for i, paragraph in enumerate(document.paragraphs):
                            if "Председатель" in paragraph.text:
                                keyword_index = i
                                break

                        new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(form_text)
                        keyword_index += 1

                        run = new_paragraph.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph.paragraph_format.line_spacing = Pt(16)
                        new_paragraph.paragraph_format.space_after = Pt(0)

                        new_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        last_index = index

                        for i, paragraph in enumerate(document.paragraphs):
                            for run in paragraph.runs:
                                if "Төраға" in run.text and run.bold:
                                    keyword_index_kz = i
                                    break

                        new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(form_text_kz)
                        keyword_index_kz += 1

                        run = new_paragraph_kz.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                        new_paragraph_kz.paragraph_format.space_after = Pt(0)

                        new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        new_paragraph.paragraph_format.keep_together = True

                        last_index_kz = index

                    bases_kz = []
                    for base in bases:
                        if base == 'представление':
                            try:
                                predstavlenie = Base.objects.get(baseName="Представление")
                                decree_list_instance.decreeBases.add(predstavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Представление не было найдено в базе данных'},
                                                    status=400)

                            bases_kz.append('ұсыныс')
                        elif base == 'рапорт':
                            try:
                                raport = Base.objects.get(baseName="Рапорт")
                                decree_list_instance.decreeBases.add(raport)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Рапорт не было найден в базе данных'}, status=400)

                            bases_kz.append('баянат')
                        elif base == 'заявление':
                            try:
                                zayavlenie = Base.objects.get(baseName="Заявление")
                                decree_list_instance.decreeBases.add(zayavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Заявление не было найдено в базе данных'}, status=400)

                            bases_kz.append('өтініш')
                        elif base == 'протокол':
                            try:
                                protocol = Base.objects.get(baseName="Протокол")
                                decree_list_instance.decreeBases.add(protocol)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Протокол не было найден в базе данных'}, status=400)
                            bases_kz.append('хаттама')
                    if len(forms) > 1:
                        # Modify bases if needed
                        modified_bases = []
                        modified_bases_kz = []
                        for base in bases:
                            if base == 'представление':
                                modified_bases.append('представления')
                                modified_bases_kz.append('ұсыныстар')
                            elif base == 'рапорт':
                                modified_bases.append('рапорта')
                                modified_bases_kz.append('баянаттар')
                            elif base == 'заявление':
                                modified_bases.append('заявления')
                                modified_bases_kz.append('өтініштер')
                            elif base == 'протокол':
                                modified_bases.append('протокола')
                                modified_bases_kz.append('хаттамалар')
                        bases = modified_bases
                        bases_kz = modified_bases_kz

                    base_text = f"\t{last_index + 1}. Настоящий приказ вступает в силу со дня подписания.\n\tОснование: {', '.join(bases)}."
                    base_text_kz = (f"\t{last_index_kz + 1}. Осы бұйрық қол қойылған күнінен бастап күшіне "
                                    f"енеді.\n\tНегіздеме: {', '.join(bases_kz)}.")

                    # Основание

                    for i, paragraph in enumerate(document.paragraphs):
                        if "Председатель" in paragraph.text:
                            keyword_index = i
                            break

                    new_paragraph = document.paragraphs[keyword_index].insert_paragraph_before(base_text)
                    keyword_index += 1

                    run = new_paragraph.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph.paragraph_format.line_spacing = Pt(16)
                    new_paragraph.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index].insert_paragraph_before('\n')

                    # Негіздеме

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Төраға" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(base_text_kz)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index_kz].insert_paragraph_before('\n')

                    doc_stream = BytesIO()
                    document.save(doc_stream)
                    doc_stream.seek(0)

                    # Prepare the HTTP response with the modified document
                    response = HttpResponse(doc_stream.read(),
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml'
                                                         '.document')
                    response['Content-Disposition'] = f'attachment; filename=Приказ о увольнении.docx'

                    doc_stream.seek(0)
                    document_id = str(uuid4())
                    document_name = f"document_{document_id}.docx"

                    minio_client = Minio(MINIO_ENDPOINT,
                                         access_key=MINIO_ACCESS_KEY,
                                         secret_key=MINIO_SECRET_KEY,
                                         secure=False)

                    minio_client.put_object(MINIO_BUCKET_NAME, document_name, data=doc_stream,
                                            length=len(doc_stream.getvalue()))
                    document_url = f"{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{document_name}"
                    print(document_url)
                    decree_list_instance.minioDocName = document_name
                    decree_list_instance.save()

                    return response


        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Упс, что-то пошло не так'}, status=405)


@csrf_exempt
def generate_komandirovka_decree(request):
    if request.method == 'POST':
        try:
            with transaction.atomic():
                body = request.body.decode('utf-8')
                data = json.loads(body)

                decreeDate = data.get('decreeDate')

                forms = data.get('forms', [])
                bases = [base['base'] for base in data.get('bases', [])]
                template_path = 'docx_generator/static/templates/komandirovka_template.docx'
                document = Document(template_path)

                keyword_index_kz = -1

                for i, paragraph in enumerate(document.paragraphs):
                    for run in paragraph.runs:
                        if "Департамент" in run.text and run.bold:
                            keyword_index_kz = i
                            break

                decree_list_instance = DecreeList.objects.create(
                    decreeType="Командировка",
                    decreeDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                    minioDocName=None
                )
                if keyword_index_kz != -1:
                    last_index = None
                    last_index_kz = None
                    departments = []
                    for index, form in enumerate(forms, start=1):
                        person_id = form.get('personId')
                        departure = form.get('departure')
                        startDate = form.get('startDate')
                        endDate = form.get('endDate')
                        choice = form.get('choice')
                        transport = form.get('transport')
                        try:
                            personInstance = Person.objects.get(pk=person_id)
                        except Person.DoesNotExist:
                            decree_list_instance.delete()
                            return JsonResponse({'error': 'Выбранного сотрудника не существует'}, status=400)

                        if KomandirovkaInfo.objects.filter(personId=personInstance, decreeId__isConfirmed=False,
                                                           decreeId__decreeType="Командировка"):
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'У сотрудника {personInstance.iin} уже существует приказ о командировке '
                                         f'который '
                                         f'не согласован'},
                                status=400)

                        departureDeparment = Department.objects.get(DepartmentNameKaz=departure)
                        personsPositionInfo = PositionInfo.objects.get(person=personInstance)

                        if departureDeparment == personsPositionInfo.department:
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'Командировка {personInstance.iin} в собственное управление невозможна'
                            },
                                status=400)
                        departments.append(personsPositionInfo.department)
                        if personInstance.inKomandirovka:
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'Сотрудник {personInstance.iin} уже находится в командировке'
                            },
                                status=400)

                        KomandirovkaInfo.objects.create(
                            startDate=startDate,
                            endDate=endDate,
                            departure=departure,
                            travelChoice=choice,
                            transport=transport,
                            personId=personInstance,
                            decreeId=decree_list_instance
                        )

                        personsPositionInfo = PositionInfo.objects.get(person=personInstance)

                        changedDepartmentNameKaz = personsPositionInfo.department.DepartmentNameKaz
                        wordsKaz = changedDepartmentNameKaz.split()
                        if wordsKaz[-1] == 'басқармасы':
                            wordsKaz[-1] = wordsKaz[-1] + 'ның'
                            changedDepartmentNameKaz = ' '.join(wordsKaz)
                        if wordsKaz[0] == 'Басқарма':
                            wordsKaz[0] = 'басқармасы' + 'ның'
                            changedDepartmentNameKaz = ' '.join(wordsKaz)

                        personsFIOKaz = personInstance.firstName + ' ' + personInstance.patronymic + ' ' + personInstance.surname

                        departureDeparment = Department.objects.get(DepartmentNameKaz=departure)
                        changedDeparture = departureDeparment.DepartmentNameKaz

                        splittedChangedDeparture = changedDeparture.split()
                        if splittedChangedDeparture[-1] == 'басқармасы':
                            changedDeparture = changedDeparture + 'на'

                        startDate = datetime.strptime(startDate, "%Y-%m-%d")
                        endDate = datetime.strptime(endDate, "%Y-%m-%d")

                        dayCount = (endDate - startDate).days

                        if startDate > endDate:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': 'Неправильно введенные даты'}, status=400)

                        startDateMonth = None
                        if startDate.month == 1:
                            startDateMonth = 'қантар'
                        if startDate.month == 2:
                            startDateMonth = 'ақпан'
                        if startDate.month == 3:
                            startDateMonth = 'наурыз'
                        if startDate.month == 4:
                            startDateMonth = 'сәуір'
                        if startDate.month == 5:
                            startDateMonth = 'мамыр'
                        if startDate.month == 6:
                            startDateMonth = 'маусым'
                        if startDate.month == 7:
                            startDateMonth = 'шілде'
                        if startDate.month == 8:
                            startDateMonth = 'тамыз'
                        if startDate.month == 9:
                            startDateMonth = 'қыркүйек'
                        if startDate.month == 10:
                            startDateMonth = 'қазан'
                        if startDate.month == 11:
                            startDateMonth = 'қараша'
                        if startDate.month == 12:
                            startDateMonth = 'желтоқсан'

                        endDateMonth = None
                        if endDate.month == 1:
                            endDateMonth = 'қантар'
                        if endDate.month == 2:
                            endDateMonth = 'ақпан'
                        if endDate.month == 3:
                            endDateMonth = 'наурыз'
                        if endDate.month == 4:
                            endDateMonth = 'сәуір'
                        if endDate.month == 5:
                            endDateMonth = 'мамыр'
                        if endDate.month == 6:
                            endDateMonth = 'маусым'
                        if endDate.month == 7:
                            endDateMonth = 'шілде'
                        if endDate.month == 8:
                            endDateMonth = 'тамыз'
                        if endDate.month == 9:
                            endDateMonth = 'қыркүйек'
                        if endDate.month == 10:
                            endDateMonth = 'қазан'
                        if endDate.month == 11:
                            endDateMonth = 'қараша'
                        if endDate.month == 12:
                            endDateMonth = 'желтоқсан'

                        if startDate.month == endDate.month:
                            dateString = str(startDate.day) + "-" + str(
                                endDate.day) + " " + startDateMonth
                        else:
                            dateString = str(
                                startDate.day) + " " + startDateMonth + " " + str(endDate.year) + " жылғы " + str(
                                endDate.day) + " " + endDateMonth

                        form_text_kz = None

                        form_text_kz = f"\t{index}. Қазақстан Республикасы Қаржылық мониторинг агенттігі (бұдан әрі - Агенттік) ________________ департаменті {changedDepartmentNameKaz} {personsPositionInfo.position.positionTitleKaz.lower()} {personsFIOKaz} {changedDeparture} {dayCount} күн мерзімге, {startDate.year} жылғы {dateString} аралығында іссапарға {choice} жіберілсін.\n\tБелгіленген жерге дейін бару және қайту жолы {transport} белгіленсін."

                        for i, paragraph in enumerate(document.paragraphs):
                            for run in paragraph.runs:
                                if "Департамент" in run.text and run.bold:
                                    keyword_index_kz = i
                                    break

                        new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(form_text_kz)
                        keyword_index_kz += 1

                        run = new_paragraph_kz.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                        new_paragraph_kz.paragraph_format.space_after = Pt(0)

                        new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        last_index_kz = index

                    bases_kz = []
                    for base in bases:
                        if base == 'представление':
                            try:
                                predstavlenie = Base.objects.get(baseName="Представление")
                                decree_list_instance.decreeBases.add(predstavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Представление не было найдено в базе данных'},
                                                    status=400)

                            bases_kz.append('ұсыныс')
                        elif base == 'рапорт':
                            try:
                                raport = Base.objects.get(baseName="Рапорт")
                                decree_list_instance.decreeBases.add(raport)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Рапорт не было найден в базе данных'}, status=400)

                            bases_kz.append('баянат')
                        elif base == 'заявление':
                            try:
                                zayavlenie = Base.objects.get(baseName="Заявление")
                                decree_list_instance.decreeBases.add(zayavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Заявление не было найдено в базе данных'}, status=400)

                            bases_kz.append('өтініш')
                        elif base == 'протокол':
                            try:
                                protocol = Base.objects.get(baseName="Протокол")
                                decree_list_instance.decreeBases.add(protocol)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Протокол не было найден в базе данных'}, status=400)
                            bases_kz.append('хаттама')
                    if len(forms) > 1:
                        # Modify bases if needed
                        modified_bases = []
                        modified_bases_kz = []
                        for base in bases:
                            if base == 'представление':
                                modified_bases.append('представления')
                                modified_bases_kz.append('ұсыныстар')
                            elif base == 'рапорт':
                                modified_bases.append('рапорта')
                                modified_bases_kz.append('баянаттар')
                            elif base == 'заявление':
                                modified_bases.append('заявления')
                                modified_bases_kz.append('өтініштер')
                            elif base == 'протокол':
                                modified_bases.append('протокола')
                                modified_bases_kz.append('хаттамалар')
                        bases = modified_bases
                        bases_kz = modified_bases_kz

                    print(departments[0].DepartmentNameKaz)
                    department_list_text = ""
                    if len(departments) == 1:
                        department_list_text = departments[0].DepartmentNameKaz

                    if len(departments) > 1:
                        for index, department in enumerate(departments):
                            department_list_text += department.DepartmentNameKaz

                            if index != len(departments) - 1:  # Check if the current department is not the last one
                                department_list_text += " және "  # Add a comma if it's not the last one

                    payment_text = f"\tАгенттіктің _________________ департаменті {department_list_text} іссапар шығындарын толық көлемде төлесін."

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Департамент" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(payment_text)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    base_text_kz = (f"\t{last_index_kz + 1}. Осы бұйрық қол қойылған күнінен бастап күшіне "
                                    f"енеді.\n\tНегіздеме: {', '.join(bases_kz)}.")

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Департамент" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(base_text_kz)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index_kz].insert_paragraph_before('\n')

                    doc_stream = BytesIO()
                    document.save(doc_stream)
                    doc_stream.seek(0)

                    # Prepare the HTTP response with the modified document
                    response = HttpResponse(doc_stream.read(),
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml'
                                                         '.document')
                    response['Content-Disposition'] = f'attachment; filename=Приказ о командировке.docx'

                    doc_stream.seek(0)
                    document_id = str(uuid4())
                    document_name = f"document_{document_id}.docx"

                    minio_client = Minio(MINIO_ENDPOINT,
                                         access_key=MINIO_ACCESS_KEY,
                                         secret_key=MINIO_SECRET_KEY,
                                         secure=False)

                    minio_client.put_object(MINIO_BUCKET_NAME, document_name, data=doc_stream,
                                            length=len(doc_stream.getvalue()))
                    document_url = f"{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{document_name}"
                    print(document_url)
                    decree_list_instance.minioDocName = document_name
                    decree_list_instance.save()

                    return response

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Упс, что-то пошло не так'}, status=405)


@csrf_exempt
def generate_otpusk_decree(request):
    if request.method == 'POST':
        try:
            with transaction.atomic():
                body = request.body.decode('utf-8')
                data = json.loads(body)

                decreeDate = data.get('decreeDate')
                forms = data.get('forms', [])
                bases = [base['base'] for base in data.get('bases', [])]

                template_path = 'docx_generator/static/templates/otpusk_template.docx'

                keyword_index_kz = -1

                document = Document(template_path)
                for i, paragraph in enumerate(document.paragraphs):
                    for run in paragraph.runs:
                        if "Департамент" in run.text and run.bold:
                            print(run.text)
                            keyword_index_kz = i
                            break

                decree_list_instance = DecreeList.objects.create(
                    decreeType="Отпуск",
                    decreeDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                    minioDocName=None
                )

                if keyword_index_kz != -1:
                    last_index = None
                    last_index_kz = None
                    template_path = None
                    for index, form in enumerate(forms, start=1):
                        person_id = form.get('personId')
                        startDate = form.get('startDate')
                        endDate = form.get('endDate')
                        otpuskType = form.get('otpuskType')
                        benefitChoice = form.get('benefitChoice')
                        priority = form.get('priority')

                        try:
                            personInstance = Person.objects.get(pk=person_id)
                        except Person.DoesNotExist:
                            decree_list_instance.delete()
                            return JsonResponse({'error': 'Выбранного сотрудника не существует'}, status=400)

                        if OtpuskInfo.objects.filter(personId=personInstance, decreeId__isConfirmed=False,
                                                     decreeId__decreeType="Отпуск"):
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'У сотрудника {personInstance.iin} уже существует приказ об отпуске '
                                         f'который '
                                         f'не согласован'},
                                status=400)

                        personsPositionInfo = PositionInfo.objects.get(person=personInstance)

                        changedDepartmentNameKaz = personsPositionInfo.department.DepartmentNameKaz
                        wordsKaz = changedDepartmentNameKaz.split()
                        if wordsKaz[-1] == 'басқармасы':
                            wordsKaz[-1] = wordsKaz[-1] + 'ның'
                            changedDepartmentNameKaz = ' '.join(wordsKaz)

                        changedPositionTitle = personsPositionInfo.position.positionTitleKaz
                        wordsKaz = changedPositionTitle.split()
                        if wordsKaz[-1] == 'уәкіл':
                            wordsKaz[-1] = wordsKaz[-1] + 'і'
                            changedPositionTitle = ' '.join(wordsKaz)

                        changedSurnameKaz = personInstance.surname

                        if personInstance.gender.genderName == 'Мужской':
                            if personInstance.surname[-2:] == 'ев' or personInstance.surname[-2:] == 'ов':
                                changedSurnameKaz = personInstance.surname + 'қа'

                        if personInstance.gender.genderName == 'Женский':
                            if personInstance.surname[-3:] == 'ева' or personInstance.surname[-3:] == 'ова':
                                changedSurnameKaz = personInstance.surname + 'ға'

                        personsFIOKaz = personInstance.firstName + ' ' + personInstance.patronymic + ' ' + changedSurnameKaz

                        startDate = datetime.strptime(startDate, "%Y-%m-%d")
                        endDate = datetime.strptime(endDate, "%Y-%m-%d")

                        if startDate >= endDate:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': f'Неправильно установленные даты'}, status=400)

                        otpuskYear = startDate.year
                        dayCount = (endDate - startDate).days + 1

                        holidays_between_dates = Holidays.objects.filter(
                            holidayDate__range=[startDate, endDate]).count()
                        dayCount = dayCount + holidays_between_dates
                        endDate = endDate + timedelta(days=holidays_between_dates)

                        print(dayCount)
                        print(holidays_between_dates)
                        vacationDays = None
                        experienced = None

                        try:
                            vacationDays = Vacation.objects.get(personId=personInstance, year=otpuskYear,
                                                                daysType="Обычные")
                        except Vacation.DoesNotExist:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': f'Сотрудник {personInstance.iin} не имеет отпускных дней'}, status=400)

                        try:
                            experienced = Vacation.objects.get(personId=personInstance, year=otpuskYear,
                                                               daysType="Стажные")
                        except Vacation.DoesNotExist:
                            print("Person doesn't have experienced vacation days")
                        # DateString
                        # (22 күнтізбелік күн,   15 күн еңбек сіңірген жылдары үшін 2 мерекелік күн)
                        # 2023 жылғы 28 желтоқсан 2024 жылғы 4 ақпан

                        dateString = None
                        startDateMonth = None
                        endDateMonth = None

                        if startDate.month == 1:
                            startDateMonth = 'қантар'
                        if startDate.month == 2:
                            startDateMonth = 'ақпан'
                        if startDate.month == 3:
                            startDateMonth = 'наурыз'
                        if startDate.month == 4:
                            startDateMonth = 'сәуір'
                        if startDate.month == 5:
                            startDateMonth = 'мамыр'
                        if startDate.month == 6:
                            startDateMonth = 'маусым'
                        if startDate.month == 7:
                            startDateMonth = 'шілде'
                        if startDate.month == 8:
                            startDateMonth = 'тамыз'
                        if startDate.month == 9:
                            startDateMonth = 'қыркүйек'
                        if startDate.month == 10:
                            startDateMonth = 'қазан'
                        if startDate.month == 11:
                            startDateMonth = 'қараша'
                        if startDate.month == 12:
                            startDateMonth = 'желтоқсан'

                        if endDate.month == 1:
                            endDateMonth = 'қантар'
                        if endDate.month == 2:
                            endDateMonth = 'ақпан'
                        if endDate.month == 3:
                            endDateMonth = 'наурыз'
                        if endDate.month == 4:
                            endDateMonth = 'сәуір'
                        if endDate.month == 5:
                            endDateMonth = 'мамыр'
                        if endDate.month == 6:
                            endDateMonth = 'маусым'
                        if endDate.month == 7:
                            endDateMonth = 'шілде'
                        if endDate.month == 8:
                            endDateMonth = 'тамыз'
                        if endDate.month == 9:
                            endDateMonth = 'қыркүйек'
                        if endDate.month == 10:
                            endDateMonth = 'қазан'
                        if endDate.month == 11:
                            endDateMonth = 'қараша'
                        if endDate.month == 12:
                            endDateMonth = 'желтоқсан'

                        oldExperiencedDaysCount = None
                        if experienced:
                            oldExperiencedDaysCount = experienced.daysCount
                        oldBasicDaysCount = vacationDays.daysCount
                        newExperiencedDaysCount = None
                        newBasicDaysCount = None

                        if experienced and experienced.daysCount != 0:
                            if vacationDays.daysCount >= dayCount - experienced.daysCount:
                                if dayCount > experienced.daysCount:
                                    if startDate.month != endDate.month:
                                        dateString = (
                                                "(" + str(dayCount - experienced.daysCount) + " күнтізбелік күн, " +
                                                str(experienced.daysCount) + " күн еңбек сіңірген жылдары үшін) " +
                                                str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                + startDateMonth + " және " + str(endDate.year) + " жылғы " + str(
                                            endDate.day) +
                                                " " + endDateMonth)
                                    else:
                                        dateString = (
                                                "(" + str(dayCount - experienced.daysCount) + " күнтізбелік күн, " +
                                                str(experienced.daysCount) + " күн еңбек сіңірген жылдары үшін) " +
                                                str(otpuskYear) + " жылғы " + str(startDate.day) + "-" + str(
                                            endDate.day) + " "
                                                + startDateMonth)
                                    newExperiencedDaysCount = 0
                                    newBasicDaysCount = oldBasicDaysCount - (dayCount - experienced.daysCount)
                                else:
                                    if priority == "Отпускные дни за выслуги лет":
                                        if startDate.month != endDate.month:
                                            dateString = ("(" + str(dayCount) + " күн еңбек сіңірген жылдары үшін) " +
                                                          str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                          + startDateMonth + " және " + str(
                                                        endDate.year) + " жылғы " + str(
                                                        endDate.day) +
                                                          " " + endDateMonth)
                                        else:
                                            dateString = ("(" + str(dayCount) + " күн еңбек сіңірген жылдары үшін) " +
                                                          str(otpuskYear) + " жылғы " + str(
                                                        startDate.day) + "-" + str(endDate.day)
                                                          + " " + startDateMonth)
                                        newExperiencedDaysCount = oldExperiencedDaysCount - dayCount
                                        newBasicDaysCount = oldBasicDaysCount

                                    if priority == "Календарные дни":
                                        if vacationDays.daysCount >= dayCount:
                                            if startDate.month != endDate.month:
                                                dateString = ("(" + str(dayCount) + " күнтізбелік күн) " +
                                                              str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                              + startDateMonth + " және " + str(
                                                            endDate.year) + " жылғы " + str(
                                                            endDate.day) +
                                                              " " + endDateMonth)
                                            else:
                                                dateString = ("(" + str(dayCount) + " күнтізбелік күн) " +
                                                              str(otpuskYear) + " жылғы " + str(
                                                            startDate.day) + "-" + str(endDate.day)
                                                              + " " + startDateMonth)
                                            newExperiencedDaysCount = oldExperiencedDaysCount
                                            newBasicDaysCount = oldBasicDaysCount - dayCount
                                        else:
                                            return JsonResponse(
                                                {
                                                    'error': f'У сотрудника {personInstance.iin} недостаточно календарных '
                                                             f'отпускных дней на {otpuskYear} '},
                                                status=400)

                            else:
                                return JsonResponse(
                                    {
                                        'error': f'У сотрудника {personInstance.iin} недостаточно отпускных дней на {otpuskYear} '
                                                 f'год, осталось обычных {vacationDays.daysCount}, '
                                                 f'стажных {experienced.daysCount}'},
                                    status=400)
                            if holidays_between_dates > 0:
                                if dayCount >= experienced.daysCount:
                                    if startDate.month != endDate.month:
                                        dateString = ("(" + str(
                                            dayCount - experienced.daysCount - holidays_between_dates) + " күнтізбелік күн, " +
                                                      str(experienced.daysCount) + " күн еңбек сіңірген жылдары үшін, " +
                                                      str(holidays_between_dates) + " мерекелік күн) " +
                                                      str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                      + startDateMonth + " және " + str(endDate.year) + " жылғы " + str(
                                                    endDate.day) +
                                                      " " + endDateMonth)
                                    else:
                                        dateString = ("(" + str(
                                            dayCount - experienced.daysCount - holidays_between_dates) + " күнтізбелік күн, " +
                                                      str(experienced.daysCount) + " күн еңбек сіңірген жылдары үшін, " +
                                                      str(holidays_between_dates) + " мерекелік күн) " +
                                                      str(otpuskYear) + " жылғы " + str(
                                                    startDate.day) + "-" + str(endDate.day)
                                                      + " " + startDateMonth)
                                    newBasicDaysCount = oldBasicDaysCount - (
                                            (dayCount - oldExperiencedDaysCount) - holidays_between_dates)
                                    newExperiencedDaysCount = 0
                                else:
                                    if priority == "Отпускные дни за выслуги лет":
                                        if startDate.month != endDate.month:
                                            dateString = ("(" +
                                                          str(dayCount - holidays_between_dates) + "күн еңбек сіңірген жылдары "
                                                                                                   "үшін, " +
                                                          str(holidays_between_dates) + " мерекелік күн) " +
                                                          str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                          + startDateMonth + " және " + str(
                                                        endDate.year) + " жылғы " + str(
                                                        endDate.day) +
                                                          " " + endDateMonth)
                                        else:
                                            dateString = ("(" +
                                                          str(dayCount - holidays_between_dates) + "күн еңбек сіңірген жылдары "
                                                                                                   "үшін, " +
                                                          str(holidays_between_dates) + " мерекелік күн) " +
                                                          str(otpuskYear) + " жылғы " + str(
                                                        startDate.day) + "-" + str(endDate.day)
                                                          + " " + startDateMonth)
                                        newExperiencedDaysCount = oldExperiencedDaysCount - (
                                                dayCount - holidays_between_dates)
                                    if priority == "Календарные дни":
                                        if vacationDays.daysCount >= dayCount:
                                            if startDate.month != endDate.month:
                                                dateString = ("(" +
                                                              str(dayCount - holidays_between_dates) + " күнтізбелік күн, " +
                                                              str(holidays_between_dates) + " мерекелік күн) " +
                                                              str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                              + startDateMonth + " және " + str(
                                                            endDate.year) + " жылғы " + str(
                                                            endDate.day) +
                                                              " " + endDateMonth)
                                            else:
                                                dateString = ("(" +
                                                              str(dayCount - holidays_between_dates) + " күнтізбелік күн, " +
                                                              str(holidays_between_dates) + " мерекелік күн) " +
                                                              str(otpuskYear) + " жылғы " + str(
                                                            startDate.day) + "-" + str(endDate.day)
                                                              + " " + startDateMonth)
                                            newBasicDaysCount = oldBasicDaysCount - (dayCount - holidays_between_dates)

                                        else:
                                            return JsonResponse(
                                                {
                                                    'error': f'У сотрудника {personInstance.iin} недостаточно календарных отпускных дней на {otpuskYear} '},
                                                status=400)

                        else:
                            if vacationDays.daysCount >= dayCount:
                                if startDate.month != endDate.month:
                                    dateString = ("(" + str(dayCount) + " күнтізбелік күн) " +
                                                  str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                  + startDateMonth + " және " + str(endDate.year) + " жылғы " + str(
                                                endDate.day) +
                                                  " " + endDateMonth)
                                else:
                                    dateString = ("(" + str(dayCount) + " күнтізбелік күн) " +
                                                  str(otpuskYear) + " жылғы " + str(startDate.day) + "-"
                                                  + str(endDate.day) + " " + str(startDateMonth))
                                newBasicDaysCount = oldBasicDaysCount - dayCount
                                newExperiencedDaysCount = oldExperiencedDaysCount

                            else:
                                return JsonResponse(
                                    {
                                        'error': f'У сотрудника {personInstance.iin} недостаточно отпускных дней на {otpuskYear} '
                                                 f'год, осталось обычных {vacationDays.daysCount}'}, status=400)
                            if holidays_between_dates > 0:
                                if startDate.month != endDate.month:
                                    dateString = ("(" + str(dayCount - holidays_between_dates) + " күнтізбелік күн, " +
                                                  str(holidays_between_dates) + " мерекелік күн) " +
                                                  str(otpuskYear) + " жылғы " + str(startDate.day) + " "
                                                  + startDateMonth + " және " + str(endDate.year) + " жылғы " + str(
                                                endDate.day) +
                                                  " " + endDateMonth)
                                else:
                                    dateString = ("(" + str(dayCount - holidays_between_dates) + " күнтізбелік күн, " +
                                                  str(holidays_between_dates) + " мерекелік күн) " +
                                                  str(otpuskYear) + " жылғы " + str(startDate.day) + "-"
                                                  + str(endDate.day) + " " + str(startDateMonth))
                                newBasicDaysCount = oldBasicDaysCount - (dayCount - holidays_between_dates)
                                newExperiencedDaysCount = oldExperiencedDaysCount

                        print("oldBasicDaysCount: ", oldBasicDaysCount)
                        print("newBasicDaysCount: ", newBasicDaysCount)

                        print("oldExperiencedDaysCount: ", oldExperiencedDaysCount)
                        print("newExperiencedDaysCount: ", newExperiencedDaysCount)

                        print(dateString)

                        form_text_kz = None

                        if otpuskType == 'Отпуск Кратко':

                            newBasicDaysCount = oldBasicDaysCount
                            newExperiencedDaysCount = oldExperiencedDaysCount

                            if startDate.month == endDate.month:
                                dateString = str(otpuskYear) + " жылғы " + str(startDate.day) + "-" + str(
                                    endDate.day) + " " + startDateMonth
                            else:
                                dateString = str(otpuskYear) + " жылғы " + str(
                                    startDate.day) + " " + startDateMonth + " " + str(endDate.year) + " жылғы " + str(
                                    endDate.day) + " " + endDateMonth

                            form_text_kz = f"\t{index}. Қазақстан Республикасы Қаржылық мониторинг агенттігі _____________________ департаменті {changedDepartmentNameKaz} {changedPositionTitle.lower()} {personsFIOKaz} қысқа мерзімді ақылы демалысы {dayCount} күн мерзімінде, {dateString} аралығында берілсін."

                        OtpuskInfo.objects.create(
                            startDate=startDate,
                            endDate=endDate,
                            otpuskType=otpuskType,
                            benefitChoice=benefitChoice,
                            oldBasicDaysCount=oldBasicDaysCount,
                            newBasicDaysCount=newBasicDaysCount,
                            oldExperiencedDaysCount=oldExperiencedDaysCount,
                            newExperiencedDaysCount=newExperiencedDaysCount,
                            personId=personInstance,
                            decreeId=decree_list_instance
                        )

                        if benefitChoice == "Без пособия":
                            form_text_kz = f"\t{index}. Қазақстан Республикасы Қаржылық мониторинг агенттігінің (бұдан әрі - Агенттік) __________________ департаменті {changedDepartmentNameKaz} {changedPositionTitle.lower()} {personsFIOKaz} {startDate.year} жылғы жыл сайынғы еңбек демалысының пайдаланбаған күндері {dayCount} күн мерзімінде, {dateString} аралығында берілсін."

                        if benefitChoice == "С пособием":
                            form_text_kz = f"\t{index}. Қазақстан Республикасы Қаржылық мониторинг агенттігінің (бұдан әрі - Агенттік) __________________ департаменті {changedDepartmentNameKaz} {changedPositionTitle.lower()} {personsFIOKaz} {startDate.year} жылғы жыл сайынғы еңбек демалысының пайдаланбаған күндері {dayCount} күн мерзімінде, {dateString} аралығында, сауықтыруға екі лауазымдық айлықақы мөлшерінде жәрдемақы төлене отырып берілсін."

                        for i, paragraph in enumerate(document.paragraphs):
                            for run in paragraph.runs:
                                if "Департамент" in run.text and run.bold:
                                    keyword_index_kz = i
                                    break

                        new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(form_text_kz)
                        keyword_index_kz += 1

                        run = new_paragraph_kz.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                        new_paragraph_kz.paragraph_format.space_after = Pt(0)

                        new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        last_index_kz = index

                    bases_kz = []
                    for base in bases:
                        if base == 'представление':
                            try:
                                predstavlenie = Base.objects.get(baseName="Представление")
                                decree_list_instance.decreeBases.add(predstavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Представление не было найдено в базе данных'},
                                                    status=400)

                            bases_kz.append('ұсыныс')
                        elif base == 'рапорт':
                            try:
                                raport = Base.objects.get(baseName="Рапорт")
                                decree_list_instance.decreeBases.add(raport)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Рапорт не было найден в базе данных'}, status=400)

                            bases_kz.append('баянат')
                        elif base == 'заявление':
                            try:
                                zayavlenie = Base.objects.get(baseName="Заявление")
                                decree_list_instance.decreeBases.add(zayavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Заявление не было найдено в базе данных'}, status=400)

                            bases_kz.append('өтініш')
                        elif base == 'протокол':
                            try:
                                protocol = Base.objects.get(baseName="Протокол")
                                decree_list_instance.decreeBases.add(protocol)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Протокол не было найден в базе данных'}, status=400)
                            bases_kz.append('хаттама')
                    if len(forms) > 1:
                        # Modify bases if needed
                        modified_bases = []
                        modified_bases_kz = []
                        for base in bases:
                            if base == 'представление':
                                modified_bases.append('представления')
                                modified_bases_kz.append('ұсыныстар')
                            elif base == 'рапорт':
                                modified_bases.append('рапорта')
                                modified_bases_kz.append('баянаттар')
                            elif base == 'заявление':
                                modified_bases.append('заявления')
                                modified_bases_kz.append('өтініштер')
                            elif base == 'протокол':
                                modified_bases.append('протокола')
                                modified_bases_kz.append('хаттамалар')
                        bases = modified_bases
                        bases_kz = modified_bases_kz

                    base_text_kz = (f"\t{last_index_kz + 1}. Осы бұйрық қол қойылған күнінен бастап күшіне "
                                    f"енеді.\n\tНегіздеме: {', '.join(bases_kz)}.")

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Департамент" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(base_text_kz)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index_kz].insert_paragraph_before('\n')

                    doc_stream = BytesIO()
                    document.save(doc_stream)
                    doc_stream.seek(0)

                    # Prepare the HTTP response with the modified document
                    response = HttpResponse(doc_stream.read(),
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml'
                                                         '.document')
                    response['Content-Disposition'] = f'attachment; filename=Приказ об отпуске.docx'

                    doc_stream.seek(0)
                    document_id = str(uuid4())
                    document_name = f"document_{document_id}.docx"

                    minio_client = Minio(MINIO_ENDPOINT,
                                         access_key=MINIO_ACCESS_KEY,
                                         secret_key=MINIO_SECRET_KEY,
                                         secure=False)

                    minio_client.put_object(MINIO_BUCKET_NAME, document_name, data=doc_stream,
                                            length=len(doc_stream.getvalue()))
                    document_url = f"{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{document_name}"
                    print(document_url)
                    decree_list_instance.minioDocName = document_name
                    decree_list_instance.save()

                    return response
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    return JsonResponse({'error': 'Упс, что-то пошло не так'}, status=405)


@csrf_exempt
def generate_otpusk_otziv_decree(request):
    if request.method == 'POST':
        try:
            with transaction.atomic():
                body = request.body.decode('utf-8')
                data = json.loads(body)

                decreeDate = data.get('decreeDate')
                forms = data.get('forms', [])
                bases = [base['base'] for base in data.get('bases', [])]

                template_path = 'docx_generator/static/templates/otpusk_otziv_template.docx'

                keyword_index_kz = -1

                document = Document(template_path)
                for i, paragraph in enumerate(document.paragraphs):
                    for run in paragraph.runs:
                        if "Департамент" in run.text and run.bold:
                            print(run.text)
                            keyword_index_kz = i
                            break

                decree_list_instance = DecreeList.objects.create(
                    decreeType="Отпуск",
                    decreeDate=datetime.strptime(decreeDate, '%Y-%m-%d').date(),
                    minioDocName=None
                )

                if keyword_index_kz != -1:
                    last_index = None
                    last_index_kz = None
                    for index, form in enumerate(forms, start=1):
                        person_id = form.get('personId')
                        otpuskType = form.get('otpuskType')
                        otzivDate = form.get('otzivDate')

                        try:
                            personInstance = Person.objects.get(pk=person_id)
                        except Person.DoesNotExist:
                            decree_list_instance.delete()
                            return JsonResponse({'error': 'Выбранного сотрудника не существует'}, status=400)

                        if OtpuskInfo.objects.filter(personId=personInstance, decreeId__isConfirmed=False,
                                                     decreeId__decreeType="Отпуск"):
                            transaction.set_rollback(True)
                            return JsonResponse({
                                'error': f'У сотрудника {personInstance.iin} уже существует приказ об отпуске '
                                         f'который '
                                         f'не согласован'},
                                status=400)

                        personsPositionInfo = PositionInfo.objects.get(person=personInstance)

                        changedDepartmentNameKaz = personsPositionInfo.department.DepartmentNameKaz
                        wordsKaz = changedDepartmentNameKaz.split()
                        if wordsKaz[-1] == 'басқармасы':
                            wordsKaz[-1] = wordsKaz[-1] + 'ның'
                            changedDepartmentNameKaz = ' '.join(wordsKaz)

                        changedPositionTitle = personsPositionInfo.position.positionTitleKaz
                        wordsKaz = changedPositionTitle.split()
                        if wordsKaz[-1] == 'уәкіл':
                            wordsKaz[-1] = wordsKaz[-1] + 'і'
                            changedPositionTitle = ' '.join(wordsKaz)

                        changedSurnameKaz = personInstance.surname

                        if personInstance.gender.genderName == 'Мужской':
                            if personInstance.surname[-2:] == 'ев' or personInstance.surname[-2:] == 'ов':
                                changedSurnameKaz = personInstance.surname

                        if personInstance.gender.genderName == 'Женский':
                            if personInstance.surname[-3:] == 'ева' or personInstance.surname[-3:] == 'ова':
                                changedSurnameKaz = personInstance.surname

                        personsFIOKaz = personInstance.firstName + ' ' + personInstance.patronymic + ' ' + changedSurnameKaz

                        if not personInstance.inVacation:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': f'Сотрудник {personInstance.iin} не находится в отпуске'},
                                                status=400)

                        otpuskInfo = OtpuskInfo.objects.filter(decreeId__decreeType="Отпуск", personId=personInstance,
                                                               decreeId__isConfirmed=True).last()

                        if not otpuskInfo:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': f'У сотрудника {personInstance.iin} нету приказа об отпуске'},
                                                status=400)

                        otzivDate = datetime.strptime(otzivDate, "%Y-%m-%d").date()
                        print(otpuskInfo.startDate)
                        print(otpuskInfo.endDate)
                        print(otzivDate)
                        if otzivDate < otpuskInfo.startDate or otzivDate > otpuskInfo.endDate:
                            transaction.set_rollback(True)
                            return JsonResponse({'error': f'Неправильно введенная дата отзыва'},
                                                status=400)


                        vacationDays = None
                        experienced = None

                        try:
                            vacationDays = Vacation.objects.get(personId=personInstance, year=otpuskInfo.startDate.year,
                                                                daysType="Обычные")
                        except Vacation.DoesNotExist:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {'error': f'Нету сущности отпускные дни для сотрудника {personInstance.iin}'},
                                status=400)

                        try:
                            experienced = Vacation.objects.get(personId=personInstance, year=otpuskInfo.startDate.year,
                                                               daysType="Стажные")
                        except Vacation.DoesNotExist:
                            print("Person doesn't have experienced vacation days")

                        monthStringKaz = None

                        if otzivDate.month == 1:
                            monthStringKaz = 'қантардан'
                        if otzivDate.month == 2:
                            monthStringKaz = 'ақпаннан'
                        if otzivDate.month == 3:
                            monthStringKaz = 'наурыздан'
                        if otzivDate.month == 4:
                            monthStringKaz = 'сәуірден'
                        if otzivDate.month == 5:
                            monthStringKaz = 'мамырдан'
                        if otzivDate.month == 6:
                            monthStringKaz = 'маусымнан'
                        if otzivDate.month == 7:
                            monthStringKaz = 'шілдеден'
                        if otzivDate.month == 8:
                            monthStringKaz = 'тамыздан'
                        if otzivDate.month == 9:
                            monthStringKaz = 'қыркүйектен'
                        if otzivDate.month == 10:
                            monthStringKaz = 'қазаннан'
                        if otzivDate.month == 11:
                            monthStringKaz = 'қарашадан'
                        if otzivDate.month == 12:
                            monthStringKaz = 'желтоқсаннан'

                        dateString = str(otzivDate.year) + " жылғы " + str(otzivDate.day) + " " + monthStringKaz

                        oldBasicDaysCount = None
                        if vacationDays is not None:
                            oldBasicDaysCount = vacationDays.daysCount

                        oldExperiencedDaysCount = None
                        if experienced is not None:
                            oldExperiencedDaysCount = experienced.daysCount

                        newBasicDaysCount = oldBasicDaysCount + (otpuskInfo.endDate - otzivDate).days + 1
                        newExperiencedDaysCount = oldExperiencedDaysCount

                        OtpuskInfo.objects.create(
                            startDate=otpuskInfo.startDate,
                            endDate=otpuskInfo.endDate,
                            otpuskType=otpuskType,
                            benefitChoice=None,
                            otzivDate=otzivDate,
                            oldBasicDaysCount=oldBasicDaysCount,
                            newBasicDaysCount=newBasicDaysCount,
                            oldExperiencedDaysCount=oldExperiencedDaysCount,
                            newExperiencedDaysCount=newExperiencedDaysCount,
                            personId=personInstance,
                            decreeId=decree_list_instance
                        )

                        form_text_kz = f"\t{index}. Қазақстан Республикасы Қаржылық мониторинг агенттігінің (бұдан әрі - Агенттік) __________________ департаменті {changedDepartmentNameKaz} {changedPositionTitle.lower()} {personsFIOKaz} жыл сайынғы еңбек демалысынан {dateString} бастап шақыртылсын."

                        for i, paragraph in enumerate(document.paragraphs):
                            for run in paragraph.runs:
                                if "Департамент" in run.text and run.bold:
                                    keyword_index_kz = i
                                    break

                        new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(form_text_kz)
                        keyword_index_kz += 1

                        run = new_paragraph_kz.runs[0]
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

                        new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                        new_paragraph_kz.paragraph_format.space_after = Pt(0)

                        new_paragraph_kz.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        last_index_kz = index

                    bases_kz = []
                    for base in bases:
                        if base == 'представление':
                            try:
                                predstavlenie = Base.objects.get(baseName="Представление")
                                decree_list_instance.decreeBases.add(predstavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Представление не было найдено в базе данных'},
                                                    status=400)

                            bases_kz.append('ұсыныс')
                        elif base == 'рапорт':
                            try:
                                raport = Base.objects.get(baseName="Рапорт")
                                decree_list_instance.decreeBases.add(raport)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Рапорт не было найден в базе данных'}, status=400)

                            bases_kz.append('баянат')
                        elif base == 'заявление':
                            try:
                                zayavlenie = Base.objects.get(baseName="Заявление")
                                decree_list_instance.decreeBases.add(zayavlenie)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Заявление не было найдено в базе данных'}, status=400)

                            bases_kz.append('өтініш')
                        elif base == 'протокол':
                            try:
                                protocol = Base.objects.get(baseName="Протокол")
                                decree_list_instance.decreeBases.add(protocol)
                            except Base.DoesNotExist:
                                decree_list_instance.delete()
                                return JsonResponse({'error': 'Протокол не было найден в базе данных'}, status=400)
                            bases_kz.append('хаттама')
                    if len(forms) > 1:
                        # Modify bases if needed
                        modified_bases = []
                        modified_bases_kz = []
                        for base in bases:
                            if base == 'представление':
                                modified_bases.append('представления')
                                modified_bases_kz.append('ұсыныстар')
                            elif base == 'рапорт':
                                modified_bases.append('рапорта')
                                modified_bases_kz.append('баянаттар')
                            elif base == 'заявление':
                                modified_bases.append('заявления')
                                modified_bases_kz.append('өтініштер')
                            elif base == 'протокол':
                                modified_bases.append('протокола')
                                modified_bases_kz.append('хаттамалар')
                        bases = modified_bases
                        bases_kz = modified_bases_kz

                    base_text_kz = (f"\t{last_index_kz + 1}. Осы бұйрық қол қойылған күнінен бастап күшіне "
                                    f"енеді.\n\tНегіздеме: {', '.join(bases_kz)}.")

                    for i, paragraph in enumerate(document.paragraphs):
                        for run in paragraph.runs:
                            if "Департамент" in run.text and run.bold:
                                keyword_index_kz = i
                                break

                    new_paragraph_kz = document.paragraphs[keyword_index_kz].insert_paragraph_before(base_text_kz)
                    keyword_index_kz += 1

                    run = new_paragraph_kz.runs[0]
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

                    new_paragraph_kz.paragraph_format.line_spacing = Pt(16)
                    new_paragraph_kz.paragraph_format.space_after = Pt(0)

                    document.paragraphs[keyword_index_kz].insert_paragraph_before('\n')

                    doc_stream = BytesIO()
                    document.save(doc_stream)
                    doc_stream.seek(0)

                    # Prepare the HTTP response with the modified document
                    response = HttpResponse(doc_stream.read(),
                                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml'
                                                         '.document')
                    response['Content-Disposition'] = f'attachment; filename=Приказ об отпуске отзыве.docx'

                    doc_stream.seek(0)
                    document_id = str(uuid4())
                    document_name = f"document_{document_id}.docx"

                    minio_client = Minio(MINIO_ENDPOINT,
                                         access_key=MINIO_ACCESS_KEY,
                                         secret_key=MINIO_SECRET_KEY,
                                         secure=False)

                    minio_client.put_object(MINIO_BUCKET_NAME, document_name, data=doc_stream,
                                            length=len(doc_stream.getvalue()))
                    document_url = f"{MINIO_ENDPOINT}/{MINIO_BUCKET_NAME}/{document_name}"
                    print(document_url)
                    decree_list_instance.minioDocName = document_name
                    decree_list_instance.save()

                    return response

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
    return JsonResponse({'error': 'Упс, что-то пошло не так'}, status=405)

