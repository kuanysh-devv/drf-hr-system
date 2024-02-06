import io
from docx.shared import Pt
from docx.shared import RGBColor
from django.http import JsonResponse, HttpResponse
from rest_framework import viewsets
from rest_framework.permissions import IsAuthenticated
from xlsxwriter import Workbook
import pandas as pd
from staffing_table.models import Vacancy, StaffingTable
from location.models import Location, Department
from location.serializers import DepartmentSerializer
from person.models import Person
from person.serializers import PersonSerializer
from position.models import PositionInfo, Position
from position.serializers import PositionSerializer
from staffing_table.models import StaffingTable
from staffing_table.serializers import StaffingTableSerializer, VacancySerializer
from django.http import HttpResponse
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from datetime import date
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from rest_framework.authentication import TokenAuthentication
from rest_framework.permissions import IsAuthenticated
from rest_framework.decorators import api_view, authentication_classes, permission_classes
from rest_framework.decorators import action


class StaffingTableViewSet(viewsets.ModelViewSet):
    queryset = StaffingTable.objects.all()
    serializer_class = StaffingTableSerializer
    permission_classes = (IsAuthenticated,)

    @action(detail=False, methods=['get'])
    def downloadStaffingTable(self, request, *args, **kwargs):
        departmentChosen = request.GET.get('department')

        doc = Document()
        doc.add_paragraph('')
        # Set the paper orientation to landscape
        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        section.left_margin = Pt(50)  # Set left margin in points
        section.right_margin = Pt(50)

        # Add the title and date
        title_paragraph = doc.add_paragraph("СПИСОК", 'Normal')
        for run in title_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        title_paragraph.paragraph_format.alignment = 1  # Center alignment
        title_paragraph.paragraph_format.line_spacing = Pt(10)  # Set line spacing

        subtitle_paragraph = doc.add_paragraph("сотрудников департамента АФМ РК на ", 'Normal')
        for run in subtitle_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        subtitle_paragraph.paragraph_format.alignment = 1  # Center alignment
        subtitle_paragraph.paragraph_format.line_spacing = Pt(10)  # Set line spacing

        # Add date with custom font and size
        formatted_date = date.today().strftime("%d.%m.%Y г.")
        run = subtitle_paragraph.add_run(formatted_date)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(13)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)  # Black color

        try:
            rukDep = Person.objects.get(positionInfo__position__positionTitle__exact="Руководитель департамента")
            zamRukDep = Person.objects.get(
                positionInfo__position__positionTitle__exact="Заместитель руководителя департамента")
        except Person.DoesNotExist:
            rukDep = None
            zamRukDep = None

        table = doc.add_table(rows=1, cols=3)
        table.style = 'TableGrid'
        header_cells = table.rows[0].cells
        header_cells[0].text = '№'
        header_cells[1].text = 'Наименование должности'
        header_cells[2].text = 'Ф.И.О. сотрудника'

        header_cells[0].width = Pt(1.0)
        header_cells[1].width = Pt(350.0)
        header_cells[2].width = Pt(600.0)

        header_row = table.rows[0]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="CCC0D9"/>'.format(nsdecls('w')))
        header_row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="CCC0D9"/>'.format(nsdecls('w')))
        header_row.cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="CCC0D9"/>'.format(nsdecls('w')))
        header_row.cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.bold = True
                paragraph.alignment = 1  # Center alignment
                paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if rukDep is not None and zamRukDep is not None:

            cells = table.add_row().cells
            cells[0].text = str(1) + '.'
            cells[1].text = rukDep.positionInfo.position.positionTitle
            full_name = f"{rukDep.firstName} {rukDep.surname} {rukDep.patronymic}"
            cells[2].text = full_name
            cells[0].width = Pt(1.0)
            cells[1].width = Pt(300.0)
            cells[2].width = Pt(500.0)

            for cell in cells:
                for run in cell.paragraphs[0].runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

            cells = table.add_row().cells
            cells[0].text = str(2) + '.'
            cells[1].text = zamRukDep.positionInfo.position.positionTitle
            full_name = f"{zamRukDep.firstName} {zamRukDep.surname} {zamRukDep.patronymic}"
            cells[2].text = full_name
            cells[0].width = Pt(1.0)
            cells[1].width = Pt(300.0)
            cells[2].width = Pt(500.0)

            for cell in cells:
                for run in cell.paragraphs[0].runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

        # Iterate over each department and create a table
        allEmployeeCount = Person.objects.all().count() - 2
        allVacanciesCount = Vacancy.objects.all().count()

        if departmentChosen == "Все управления":
            for department in Department.objects.all():
                # Add a row for the department name
                row = table.add_row()
                row.cells[0].text = department.DepartmentName
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                            run.font.bold = True
                        paragraph.alignment = 1  # Center alignment
                        paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row.cells[0].merge(row.cells[-1])

                # Get positions in the current department from StaffingTable
                positions_in_department = StaffingTable.objects.filter(
                    staffing_table_department=department).values_list(
                    'staffing_table_position', flat=True).order_by(
                    '-staffing_table_position__order')

                employee_count = Person.objects.filter(positionInfo__department=department).count()

                # Get the count of available vacancies for the current department
                vacancy_count = Vacancy.objects.filter(department=department).count()

                # Determine the maximum count between employees and vacancies
                max_count = employee_count + vacancy_count
                # Iterate through positions in the department
                num = 1
                for position_id in positions_in_department:
                    position = Position.objects.get(pk=position_id)

                    # Filter employees for the current department and position
                    employees_in_department = Person.objects.filter(positionInfo__department=department,
                                                                    positionInfo__position=position)
                    # Filter vacancies for the current department and position
                    vacancies_for_position = Vacancy.objects.filter(department=department, position=position)

                    # Add employee data to the table
                    for employee in employees_in_department:

                        cells = table.add_row().cells
                        cells[0].text = str(num) + '.'
                        cells[1].text = employee.positionInfo.position.positionTitle
                        full_name = f"{employee.firstName} {employee.surname} {employee.patronymic}"
                        cells[2].text = full_name
                        cells[0].width = Pt(1.0)
                        cells[1].width = Pt(300.0)
                        cells[2].width = Pt(500.0)

                        for cell in cells:
                            for run in cell.paragraphs[0].runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(13)
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                        num = num + 1
                    # Add vacancy data to the table
                    for vacancy in vacancies_for_position:

                        cells = table.add_row().cells
                        cells[0].text = str(num) + '.'
                        cells[1].text = vacancy.position.positionTitle
                        # Format the date in the desired format
                        formatted_date = "Вакансия " + vacancy.available_date.strftime("%d.%m.%Y г.")
                        cells[2].text = formatted_date
                        cells[0].width = Pt(1.0)
                        cells[1].width = Pt(300.0)
                        cells[2].width = Pt(500.0)

                        for cell in cells:
                            for run in cell.paragraphs[0].runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(13)
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                                run.font.italic = True
                                run.font.bold = True

                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w')))
                        cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w')))
                        cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w')))
                        cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
                        num = num + 1
            row = table.add_row()
            row.cells[0].text = "Всего по департаменту"
            row.cells[2].text = str(allEmployeeCount + allVacanciesCount)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            row.cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(13)
                        run.font.bold = True
                    paragraph.alignment = 1  # Center alignment
                    paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[0].merge(row.cells[1])
        else:
            department = Department.objects.get(DepartmentName=departmentChosen)
            allEmployeeCount = Person.objects.filter(positionInfo__department=department).count()
            allVacanciesCount = Vacancy.objects.filter(department=department).count()
            row = table.add_row()
            row.cells[0].text = department.DepartmentName
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
            row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(13)
                        run.font.bold = True
                    paragraph.alignment = 1  # Center alignment
                    paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[0].merge(row.cells[-1])

            # Get positions in the current department from StaffingTable
            positions_in_department = StaffingTable.objects.filter(staffing_table_department=department).values_list(
                'staffing_table_position', flat=True).order_by(
                '-staffing_table_position__order')

            employee_count = Person.objects.filter(positionInfo__department=department).count()

            # Get the count of available vacancies for the current department
            vacancy_count = Vacancy.objects.filter(department=department).count()

            # Determine the maximum count between employees and vacancies
            max_count = employee_count + vacancy_count
            # Iterate through positions in the department
            num = 1
            for position_id in positions_in_department:
                position = Position.objects.get(pk=position_id)

                # Filter employees for the current department and position
                employees_in_department = Person.objects.filter(positionInfo__department=department,
                                                                positionInfo__position=position)
                # Filter vacancies for the current department and position
                vacancies_for_position = Vacancy.objects.filter(department=department, position=position)

                # Add employee data to the table
                for employee in employees_in_department:

                    cells = table.add_row().cells
                    cells[0].text = str(num) + '.'
                    cells[1].text = employee.positionInfo.position.positionTitle
                    full_name = f"{employee.firstName} {employee.surname} {employee.patronymic}"
                    cells[2].text = full_name
                    cells[0].width = Pt(1.0)
                    cells[1].width = Pt(300.0)
                    cells[2].width = Pt(500.0)

                    for cell in cells:
                        for run in cell.paragraphs[0].runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                    num = num + 1
                # Add vacancy data to the table
                for vacancy in vacancies_for_position:

                    cells = table.add_row().cells
                    cells[0].text = str(num) + '.'
                    cells[1].text = vacancy.position.positionTitle
                    # Format the date in the desired format
                    formatted_date = "Вакансия " + vacancy.available_date.strftime("%d.%m.%Y г.")
                    cells[2].text = formatted_date
                    cells[0].width = Pt(1.0)
                    cells[1].width = Pt(300.0)
                    cells[2].width = Pt(500.0)

                    for cell in cells:
                        for run in cell.paragraphs[0].runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                            run.font.italic = True
                            run.font.bold = True

                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w')))
                    cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w')))
                    cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w')))
                    cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
                    num = num + 1
            row = table.add_row()
            row.cells[0].text = "Всего по управлению"
            row.cells[2].text = str(allEmployeeCount + allVacanciesCount)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            row.cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(13)
                        run.font.bold = True
                    paragraph.alignment = 1  # Center alignment
                    paragraph.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[0].merge(row.cells[1])

        # Create a response with the Word document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=staffing_table.docx'
        doc.save(response)

        return response

    @action(detail=False, methods=['get'])
    def getStaffingTable(self, request, *args, **kwargs):
        department_id = request.GET.get('department_id')

        try:
            department = Department.objects.get(pk=department_id)

            StaffingTableInstances = StaffingTable.objects.filter(staffing_table_department=department).distinct(
                'staffing_table_position')
            # GlavExpertDc
            # ExpertDC
            positionList = []
            for staffTable in StaffingTableInstances:
                position_data = PositionSerializer(staffTable.staffing_table_position).data
                currentPositionInfos = PositionInfo.objects.filter(position=staffTable.staffing_table_position,
                                                                   department=staffTable.staffing_table_department)
                personsOnPosition = Person.objects.filter(positionInfo__in=currentPositionInfos, isFired=False)
                position_data['persons'] = PersonSerializer(personsOnPosition, many=True).data

                staffing_table_entry = StaffingTable.objects.filter(
                    staffing_table_position=staffTable.staffing_table_position,
                    staffing_table_department=department
                ).first()

                if staffing_table_entry:
                    vacancies = Vacancy.objects.filter(department=department,
                                                       position=staffTable.staffing_table_position)
                    position_data['vacancies'] = VacancySerializer(vacancies, many=True).data
                positionList.append(position_data)

            departamentSerialized = DepartmentSerializer(department).data
            departamentSerialized['positionList'] = positionList
            data = departamentSerialized

            return JsonResponse(data)

        except Department.DoesNotExist:
            return JsonResponse({'error': 'Управление не найдено'}, status=404)
